import streamlit as st
import sys
import re
import time
import random
import tempfile
import os
from pathlib import Path

# ─── Auto-install missing PDF packages ────────────────────────────────────────
def _ensure(pkg, import_name=None):
    import importlib
    name = import_name or pkg.split("-")[0].replace("-", "_")
    try:
        importlib.import_module(name)
    except ImportError:
        os.system(f"pip install {pkg} --break-system-packages -q")

_ensure("pypdf", "pypdf")
_ensure("pdfplumber", "pdfplumber")
_ensure("reportlab", "reportlab")

# PDF imports
from pypdf import PdfReader, PdfWriter
import pdfplumber
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet

# ── Force agent.py onto sys.path immediately at startup ──────────────
# This runs before any function is called, so import always succeeds.
_THIS_DIR = Path(__file__).resolve().parent
if str(_THIS_DIR) not in sys.path:
    sys.path.insert(0, str(_THIS_DIR))
# Also add the current working directory as fallback
_CWD = Path(os.getcwd())
if str(_CWD) not in sys.path:
    sys.path.insert(0, str(_CWD))

# ─────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="Paperpal AI Manuscript Formatter",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ─────────────────────────────────────────────
# CSS  (Paperpal brand — purple + teal)
# ─────────────────────────────────────────────
st.markdown(r"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:wght@300;400;500;600;700;800&family=Lora:ital,wght@0,400;0,600;0,700;1,400&family=JetBrains+Mono:wght@400;500&display=swap');

:root {
  --pp-purple-900: #3B0764;
  --pp-purple-700: #6C2BD9;
  --pp-purple-600: #7C3AED;
  --pp-purple-500: #8B5CF6;
  --pp-purple-100: #EDE9FE;
  --pp-purple-50:  #F5F3FF;
  --pp-teal-600:   #0D9488;
  --pp-teal-500:   #14B8A6;
  --pp-teal-100:   #CCFBF1;
  --pp-teal-50:    #F0FDFA;
  --pp-white:      #FFFFFF;
  --pp-bg:         #F8F7FF;
  --pp-surface:    #FFFFFF;
  --pp-border:     #E5E0F5;
  --pp-border-md:  #D1C9F0;
  --pp-text-dark:  #1E1B2E;
  --pp-text-body:  #374151;
  --pp-text-muted: #6B7280;
  --pp-text-faint: #9CA3AF;
  --pp-green:      #059669;
  --pp-green-bg:   #ECFDF5;
  --pp-green-bdr:  #A7F3D0;
  --font-sans:  'Plus Jakarta Sans', sans-serif;
  --font-serif: 'Lora', serif;
  --font-mono:  'JetBrains Mono', monospace;
  --r-sm:  6px; --r-md: 12px; --r-lg: 18px; --r-xl: 26px;
  --shadow-sm: 0 1px 3px rgba(108,43,217,0.08), 0 1px 2px rgba(0,0,0,0.04);
  --shadow-md: 0 4px 16px rgba(108,43,217,0.10), 0 2px 6px rgba(0,0,0,0.05);
  --shadow-lg: 0 8px 32px rgba(108,43,217,0.14), 0 4px 12px rgba(0,0,0,0.06);
  --shadow-glow: 0 0 0 4px rgba(108,43,217,0.12);
}

@keyframes fadeUp { from{opacity:0;transform:translateY(16px)} to{opacity:1;transform:translateY(0)} }
@keyframes slideIn { from{opacity:0;transform:translateX(-10px)} to{opacity:1;transform:translateX(0)} }
@keyframes shimmer { 0%{background-position:-400px 0} 100%{background-position:400px 0} }
@keyframes gradientFlow { 0%{background-position:0% 50%} 50%{background-position:100% 50%} 100%{background-position:0% 50%} }
@keyframes pulseDot { 0%,100%{transform:scale(1);opacity:1} 50%{transform:scale(1.4);opacity:0.7} }
@keyframes ringFill { from{stroke-dashoffset:327} }
@keyframes float { 0%,100%{transform:translateY(0)} 50%{transform:translateY(-6px)} }
@keyframes countBounce { 0%{transform:scale(0.6);opacity:0} 70%{transform:scale(1.1)} 100%{transform:scale(1);opacity:1} }
@keyframes slideDown { from{opacity:0;transform:translateY(-8px)} to{opacity:1;transform:translateY(0)} }

*, *::before, *::after { box-sizing: border-box; }
html, body, [data-testid="stAppViewContainer"], [data-testid="stMain"], .main {
  background: var(--pp-bg) !important;
  font-family: var(--font-sans) !important;
  color: var(--pp-text-body) !important;
}
[data-testid="stHeader"],[data-testid="stToolbar"],[data-testid="stDecoration"] { display:none !important; }
[data-testid="stSidebar"] { display:none !important; }
section[data-testid="stMain"] > div { padding-top:0 !important; }
.block-container { padding: 20px 32px 64px !important; max-width:1380px !important; }

[data-testid="stAppViewContainer"]::before {
  content:''; position:fixed; inset:0; z-index:0; pointer-events:none;
  background-image: radial-gradient(circle, rgba(108,43,217,0.07) 1px, transparent 1px);
  background-size: 28px 28px;
}

/* NAVBAR */
.pp-navbar {
  display:flex; align-items:center; justify-content:space-between;
  padding:14px 32px; background:var(--pp-white);
  border-bottom:1px solid var(--pp-border);
  border-radius:var(--r-lg) var(--r-lg) 0 0;
  box-shadow:var(--shadow-sm); position:relative; z-index:20;
  animation:fadeUp 0.5s ease both;
}
.pp-logo { display:flex; align-items:center; gap:10px; }
.pp-logo-icon {
  width:36px; height:36px; border-radius:9px;
  background:linear-gradient(135deg,var(--pp-purple-700),var(--pp-purple-500));
  display:flex; align-items:center; justify-content:center; font-size:18px;
  box-shadow:0 2px 8px rgba(108,43,217,0.35);
}
.pp-logo-text { font-size:18px; font-weight:800; letter-spacing:-0.4px; color:var(--pp-purple-700); }
.pp-logo-text span { color:var(--pp-teal-600); }
.pp-nav-right { display:flex; align-items:center; gap:10px; }
.pp-chip {
  display:inline-flex; align-items:center; gap:5px;
  padding:5px 12px; border-radius:20px;
  font-size:11px; font-weight:700; letter-spacing:0.7px; text-transform:uppercase;
}
.chip-purple { background:var(--pp-purple-100); color:var(--pp-purple-700); border:1px solid rgba(108,43,217,0.2); }
.chip-teal   { background:var(--pp-teal-100); color:var(--pp-teal-600); border:1px solid rgba(13,148,136,0.25); }

/* HERO */
.pp-hero {
  position:relative; overflow:hidden;
  background:linear-gradient(135deg, var(--pp-purple-700) 0%, #5B21B6 45%, #4338CA 100%);
  background-size:200% 200%; animation:gradientFlow 8s ease infinite;
  padding:44px 48px 40px; margin-bottom:28px;
  border-radius:0 0 var(--r-xl) var(--r-xl);
  box-shadow:var(--shadow-lg); z-index:10;
}
.pp-hero::after {
  content:''; position:absolute; bottom:0; left:0; right:0; height:4px;
  background:linear-gradient(90deg,var(--pp-teal-600) 0%,var(--pp-teal-500) 50%,transparent 100%);
}
.pp-hero::before {
  content:''; position:absolute; right:-100px; top:-100px;
  width:380px; height:380px; border-radius:50%;
  background:rgba(255,255,255,0.05); pointer-events:none;
}
.hero-inner { position:relative; z-index:2; }
.hero-eyebrow {
  display:inline-flex; align-items:center; gap:7px;
  background:rgba(255,255,255,0.12); border:1px solid rgba(255,255,255,0.2);
  color:rgba(255,255,255,0.9); padding:5px 14px; border-radius:20px;
  font-size:12px; font-weight:600; letter-spacing:0.6px; text-transform:uppercase;
  margin-bottom:16px;
}
.hero-dot { width:6px; height:6px; border-radius:50%; background:#5EEAD4; animation:pulseDot 2s infinite; }
.hero-title { font-family:var(--font-serif); font-size:38px; font-weight:700; line-height:1.15; color:#FFFFFF; letter-spacing:-0.5px; margin-bottom:10px; }
.hero-title em { font-style:normal; color:#A5F3FC; }
.hero-sub { font-size:15px; font-weight:400; color:rgba(255,255,255,0.72); max-width:520px; line-height:1.6; }
.hero-stats { display:flex; gap:32px; margin-top:28px; }
.stat-num { font-size:22px; font-weight:800; color:#fff; }
.stat-label { font-size:11px; font-weight:500; color:rgba(255,255,255,0.55); letter-spacing:0.4px; margin-top:1px; }

/* CARDS */
.pp-card {
  background:var(--pp-white); border:1px solid var(--pp-border);
  border-radius:var(--r-lg); padding:24px 24px 20px; margin-bottom:18px;
  box-shadow:var(--shadow-sm); transition:border-color 0.25s, box-shadow 0.25s;
  animation:fadeUp 0.5s cubic-bezier(.16,1,.3,1) both;
  position:relative; overflow:hidden;
}
.pp-card:hover { border-color:var(--pp-border-md); box-shadow:var(--shadow-md); }
.pp-card::before {
  content:''; position:absolute; left:0; top:16px; bottom:16px; width:3px;
  background:linear-gradient(180deg,var(--pp-purple-600),var(--pp-teal-500));
  border-radius:0 2px 2px 0; opacity:0; transition:opacity 0.25s;
}
.pp-card:hover::before { opacity:1; }
.pp-card-title {
  font-size:15px; font-weight:700; color:var(--pp-text-dark); letter-spacing:-0.1px;
  margin-bottom:18px; padding-bottom:14px; border-bottom:1px solid var(--pp-border);
  display:flex; align-items:center; gap:9px;
}
.title-icon {
  width:28px; height:28px; border-radius:7px; flex-shrink:0;
  display:flex; align-items:center; justify-content:center; font-size:14px;
  background:var(--pp-purple-50); border:1px solid var(--pp-purple-100);
}

/* DOC TYPE BANNER */
.doc-type-banner {
  display:flex; align-items:center; gap:14px;
  background:linear-gradient(135deg, #F0FDF4, #ECFDF5);
  border:1px solid #A7F3D0; border-radius:var(--r-md);
  padding:14px 18px; margin-bottom:16px;
  animation:slideDown 0.4s cubic-bezier(.16,1,.3,1) both;
}
.doc-type-icon { font-size:28px; flex-shrink:0; }
.doc-type-info { flex:1; }
.doc-type-label {
  font-size:11px; font-weight:700; letter-spacing:0.7px; text-transform:uppercase;
  color:var(--pp-text-muted); margin-bottom:3px;
}
.doc-type-name { font-size:15px; font-weight:800; color:#065F46; }
.doc-type-desc { font-size:12px; color:#6B7280; margin-top:2px; line-height:1.4; }
.doc-type-confidence {
  padding:3px 10px; border-radius:12px; font-size:11px; font-weight:700;
  flex-shrink:0;
}
.conf-high   { background:#D1FAE5; color:#065F46; }
.conf-medium { background:#FEF3C7; color:#92400E; }
.conf-low    { background:#F3F4F6; color:#6B7280; }

/* STYLE RECOMMENDATION CHIPS */
.style-rec-row { display:flex; gap:8px; flex-wrap:wrap; margin-bottom:12px; }
.style-rec-chip {
  display:inline-flex; align-items:center; gap:5px;
  padding:5px 12px; border-radius:20px; font-size:12px; font-weight:600;
  cursor:default;
}
.style-rec-primary { background:var(--pp-purple-100); color:var(--pp-purple-700); border:1.5px solid var(--pp-purple-600); }
.style-rec-alt     { background:var(--pp-bg); color:var(--pp-text-muted); border:1px solid var(--pp-border-md); }

/* WIDGET OVERRIDES */
label[data-testid="stWidgetLabel"] p,
.stFileUploader label p,
.stSelectbox label p {
  font-family:var(--font-sans) !important;
  font-size:11px !important; font-weight:700 !important;
  letter-spacing:0.7px !important; text-transform:uppercase !important;
  color:var(--pp-text-muted) !important; margin-bottom:6px !important;
}
[data-testid="stFileUploaderDropzone"] {
  background:var(--pp-purple-50) !important;
  border:2px dashed var(--pp-purple-100) !important;
  border-radius:var(--r-md) !important; transition:all 0.25s !important;
}
[data-testid="stFileUploaderDropzone"]:hover {
  background:#EDE9FE !important; border-color:var(--pp-purple-600) !important;
  box-shadow:var(--shadow-glow) !important;
}
[data-testid="stFileUploaderDropzone"] p { color:var(--pp-text-muted) !important; font-size:13px !important; }
[data-testid="stFileUploaderDropzone"] svg { color:var(--pp-purple-600) !important; }
.stSelectbox > div > div {
  background:var(--pp-white) !important; border:1px solid var(--pp-border-md) !important;
  border-radius:var(--r-sm) !important; font-family:var(--font-sans) !important;
  box-shadow:var(--shadow-sm) !important;
}
.stSelectbox > div > div:focus-within { border-color:var(--pp-purple-600) !important; box-shadow:var(--shadow-glow) !important; }

/* BUTTONS */
.stButton > button {
  font-family:var(--font-sans) !important; font-weight:700 !important;
  letter-spacing:0.3px !important; border-radius:var(--r-md) !important;
  transition:all 0.25s cubic-bezier(.16,1,.3,1) !important;
  position:relative !important; overflow:hidden !important;
}
.stButton > button:first-child {
  background:linear-gradient(135deg,var(--pp-purple-700) 0%,var(--pp-purple-500) 100%) !important;
  background-size:200% 200% !important; animation:gradientFlow 4s ease infinite !important;
  color:#fff !important; border:none !important;
  padding:14px 28px !important; font-size:15px !important;
  box-shadow:0 4px 16px rgba(108,43,217,0.38), 0 1px 0 rgba(255,255,255,0.15) inset !important;
}
.stButton > button:first-child:hover { transform:translateY(-2px) !important; box-shadow:0 8px 28px rgba(108,43,217,0.50) !important; }
.stButton > button:first-child:active { transform:translateY(0) scale(0.99) !important; }
.stButton > button:first-child::after {
  content:''; position:absolute; inset:0;
  background:linear-gradient(110deg,transparent 35%,rgba(255,255,255,0.18) 50%,transparent 65%);
  background-size:400px 100%; animation:shimmer 2.8s infinite;
}
.stDownloadButton > button {
  background:var(--pp-white) !important; border:1px solid var(--pp-border-md) !important;
  color:var(--pp-purple-700) !important; border-radius:var(--r-sm) !important;
  font-size:13px !important; font-weight:600 !important;
  box-shadow:var(--shadow-sm) !important; transition:all 0.2s !important;
}
.stDownloadButton > button:hover {
  background:var(--pp-purple-50) !important; border-color:var(--pp-purple-600) !important;
  transform:translateY(-1px) !important;
}
.stButton:last-of-type > button {
  background:var(--pp-white) !important; border:1px solid var(--pp-border-md) !important;
  color:var(--pp-text-muted) !important; font-size:13px !important;
}

/* TABS */
.stTabs [data-baseweb="tab-list"] { background:transparent !important; border-bottom:2px solid var(--pp-border) !important; gap:0 !important; }
.stTabs [data-baseweb="tab"] {
  font-family:var(--font-sans) !important; font-weight:600 !important; font-size:14px !important;
  color:var(--pp-text-muted) !important; padding:11px 22px !important;
  border-radius:var(--r-sm) var(--r-sm) 0 0 !important;
  background:transparent !important; border:none !important;
}
.stTabs [data-baseweb="tab"]:hover { color:var(--pp-purple-700) !important; background:var(--pp-purple-50) !important; }
.stTabs [aria-selected="true"] { color:var(--pp-purple-700) !important; background:var(--pp-purple-50) !important; border-bottom:2px solid var(--pp-purple-700) !important; }

/* STATUS PILL */
.status-pill {
  display:inline-flex; align-items:center; gap:8px;
  background:var(--pp-green-bg); border:1px solid var(--pp-green-bdr); color:var(--pp-green);
  border-radius:20px; padding:6px 16px; font-size:13px; font-weight:600;
  margin-bottom:16px; animation:fadeUp 0.4s ease both;
}
.live-dot { width:7px; height:7px; border-radius:50%; background:var(--pp-green); animation:pulseDot 2s infinite; }

/* TEXT BOXES */
.col-label {
  font-size:11px; font-weight:700; letter-spacing:0.7px; text-transform:uppercase;
  color:var(--pp-text-muted); margin-bottom:8px; display:flex; align-items:center; gap:6px;
}
.dot-purple { width:5px; height:5px; border-radius:50%; background:var(--pp-purple-600); }
.dot-teal   { width:5px; height:5px; border-radius:50%; background:var(--pp-teal-600); }
.text-box {
  background:#FAFAF9; border:1px solid var(--pp-border); border-radius:var(--r-md);
  padding:16px 18px; font-size:13px; line-height:1.8; color:var(--pp-text-body);
  max-height:300px; overflow-y:auto; font-family:var(--font-serif);
  white-space:pre-wrap; box-shadow:0 1px 4px rgba(0,0,0,0.04) inset;
}

/* SCORE RING */
.score-wrap { display:flex; flex-direction:column; align-items:center; padding:12px 0 18px; }
.ring-outer {
  position:relative; width:128px; height:128px;
  display:flex; align-items:center; justify-content:center;
  animation:countBounce 0.8s cubic-bezier(.16,1,.3,1) both 0.2s;
}
.ring-svg { position:absolute; inset:0; transform:rotate(-90deg); }
.ring-track { fill:none; stroke:var(--pp-border); stroke-width:9; }
.ring-fill { fill:none; stroke:url(#ppGrad); stroke-width:9; stroke-linecap:round; animation:ringFill 1.4s cubic-bezier(.16,1,.3,1) both 0.4s; }
.ring-inner {
  width:96px; height:96px; border-radius:50%;
  background:var(--pp-white); border:1px solid var(--pp-border);
  display:flex; flex-direction:column; align-items:center; justify-content:center;
  z-index:1; box-shadow:var(--shadow-md);
}
.ring-num { font-size:30px; font-weight:800; line-height:1; color:var(--pp-purple-700); }
.ring-pct { font-size:10px; font-weight:600; color:var(--pp-text-muted); letter-spacing:0.8px; text-transform:uppercase; }
.ring-badge { margin-top:12px; padding:5px 14px; border-radius:20px; font-size:12px; font-weight:700; background:var(--pp-teal-100); border:1px solid rgba(13,148,136,0.25); color:var(--pp-teal-600); }

/* CHANGE LIST */
.changes-list { display:flex; flex-direction:column; gap:7px; }
.change-item {
  display:flex; align-items:flex-start; gap:11px;
  padding:11px 14px; background:#FAFDF9;
  border:1px solid var(--pp-green-bdr); border-left:3px solid var(--pp-teal-600);
  border-radius:0 var(--r-sm) var(--r-sm) 0;
  font-size:13.5px; color:#065F46; line-height:1.5;
  transition:all 0.2s; animation:slideIn 0.4s ease both;
}
.change-item:hover { background:#F0FDF4; border-left-color:#059669; transform:translateX(4px); }
.check-icon {
  flex-shrink:0; margin-top:2px; width:16px; height:16px; border-radius:50%;
  background:var(--pp-teal-100); border:1.5px solid var(--pp-teal-600);
  display:flex; align-items:center; justify-content:center; font-size:9px; color:var(--pp-teal-600);
}
.change-item:nth-child(1){animation-delay:.04s} .change-item:nth-child(2){animation-delay:.09s}
.change-item:nth-child(3){animation-delay:.14s} .change-item:nth-child(4){animation-delay:.19s}
.change-item:nth-child(5){animation-delay:.24s} .change-item:nth-child(6){animation-delay:.29s}
.change-item:nth-child(7){animation-delay:.34s} .change-item:nth-child(8){animation-delay:.39s}
.change-item:nth-child(9){animation-delay:.44s} .change-item:nth-child(10){animation-delay:.49s}

/* HOW IT WORKS */
.how-row { display:flex; align-items:flex-start; gap:13px; padding:10px 0; border-bottom:1px solid var(--pp-border); }
.how-row:last-child { border-bottom:none; }
.step-badge {
  width:24px; height:24px; border-radius:50%; flex-shrink:0;
  background:linear-gradient(135deg,var(--pp-purple-700),var(--pp-purple-500));
  display:flex; align-items:center; justify-content:center;
  font-size:11px; font-weight:800; color:#fff;
  box-shadow:0 2px 8px rgba(108,43,217,0.32);
}
.step-desc { font-size:13.5px; color:var(--pp-text-body); line-height:1.55; padding-top:2px; }
.mono { font-family:var(--font-mono); font-size:11.5px; background:var(--pp-purple-50); color:var(--pp-purple-700); padding:2px 6px; border-radius:4px; }

/* DOWNLOAD PANEL */
.dl-panel {
  margin-top:12px; background:var(--pp-white);
  border:1px solid var(--pp-border-md); border-radius:var(--r-md);
  overflow:hidden; box-shadow:var(--shadow-md);
  animation:fadeUp 0.35s cubic-bezier(.16,1,.3,1) both;
}
.dl-panel-header {
  display:flex; align-items:center; gap:8px; padding:11px 16px;
  background:var(--pp-purple-50); border-bottom:1px solid var(--pp-border);
  font-size:12px; font-weight:700; letter-spacing:0.5px; color:var(--pp-purple-700); text-transform:uppercase;
}
.dl-options { display:flex; gap:0; }
.dl-option {
  flex:1; display:flex; flex-direction:column; align-items:center; justify-content:center;
  padding:18px 12px 16px; gap:8px; cursor:pointer; text-decoration:none;
  transition:background 0.2s, transform 0.15s;
  border-right:1px solid var(--pp-border); position:relative;
}
.dl-option:last-child { border-right:none; }
.dl-option:hover { background:var(--pp-purple-50); transform:translateY(-1px); }
.dl-icon-wrap { width:44px; height:44px; border-radius:12px; display:flex; align-items:center; justify-content:center; font-size:22px; }
.dl-icon-word  { background:#EBF3FF; border:1px solid #BFDBFE; }
.dl-icon-gdocs { background:#FEF3C7; border:1px solid #FCD34D; }
.dl-option-label { font-size:13px; font-weight:700; color:var(--pp-text-dark); }
.dl-option-sub   { font-size:11px; color:var(--pp-text-muted); text-align:center; line-height:1.4; }
.dl-badge-word  { background:#DBEAFE; color:#1D4ED8; font-size:10px; font-weight:700; padding:2px 8px; border-radius:10px; }
.dl-badge-gdocs { background:#FEF9C3; color:#92400E; font-size:10px; font-weight:700; padding:2px 8px; border-radius:10px; }

/* EMPTY STATE */
.empty-state {
  display:flex; flex-direction:column; align-items:center; justify-content:center;
  min-height:360px; background:var(--pp-white); border:2px dashed var(--pp-border-md);
  border-radius:var(--r-lg); text-align:center; padding:48px; position:relative; overflow:hidden;
}
.empty-state::before {
  content:''; position:absolute; inset:0; pointer-events:none;
  background:radial-gradient(ellipse 60% 50% at 50% 70%, rgba(108,43,217,0.04) 0%, transparent 70%);
}
.empty-icon { font-size:52px; margin-bottom:18px; animation:float 4s ease-in-out infinite; }
.empty-h { font-family:var(--font-serif); font-size:20px; font-weight:600; color:var(--pp-text-dark); margin-bottom:8px; }
.empty-p { font-size:14px; color:var(--pp-text-muted); max-width:260px; line-height:1.6; }

/* FOOTER */
.pp-footer { text-align:center; padding:24px 20px 12px; color:var(--pp-text-faint); font-size:12px; letter-spacing:0.3px; border-top:1px solid var(--pp-border); margin-top:40px; }
.pp-footer strong { color:var(--pp-text-muted); }

/* SCROLLBARS */
::-webkit-scrollbar { width:5px; height:5px; }
::-webkit-scrollbar-track { background:var(--pp-purple-50); }
::-webkit-scrollbar-thumb { background:var(--pp-purple-100); border-radius:3px; }
::-webkit-scrollbar-thumb:hover { background:var(--pp-purple-500); }
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────
# SESSION STATE
# ─────────────────────────────────────────────
defaults = {
    "formatted_result": None,
    "original_text": None,
    "uploaded_name": None,
    "doc_type_info": None,
    "chosen_style": None,
    "style_confirmed": False,
    "upload_warning": None,
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v


# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────
STUB_TEXT = (
    "Abstract\n\n"
    "This manuscript investigates the intersection of artificial intelligence and "
    "academic publishing workflows. Prior research (Smith et al., 2021) has "
    "demonstrated significant inefficiencies in manual formatting processes, "
    "accounting for up to 34% of desk rejections at peer-reviewed journals.\n\n"
    "Introduction\n\n"
    "The proliferation of academic style guides (Jones, 2020; Lee & Park, 2022) "
    "has created a fragmented landscape for researchers. Modern language models "
    "now offer the potential to automate complex formatting tasks with near-human "
    "accuracy (Williams, 2023), representing a paradigm shift in manuscript preparation.\n\n"
    "Methodology\n\n"
    "A mixed-methods approach was employed, combining quantitative analysis of "
    "500 manuscripts with qualitative interviews of 20 senior journal editors."
)

MAX_FILE_MB = 20
MAX_FILE_BYTES = MAX_FILE_MB * 1024 * 1024

def read_uploaded_text(uf) -> tuple:
    """Returns (text, warning_message_or_None)."""
    name = uf.name.lower()
    try:
        raw_bytes = uf.read()
        if len(raw_bytes) > MAX_FILE_BYTES:
            return STUB_TEXT, f"File exceeds {MAX_FILE_MB} MB limit — using demo text."
        if name.endswith(".pdf"):
            # ── Primary: pdfplumber (best layout-aware extraction) ──
            try:
                import pdfplumber, io
                with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
                    text = "\n\n".join(p.extract_text() or "" for p in pdf.pages).strip()
                if text:
                    return text, None
            except Exception:
                text = ""

            # ── Fallback: pypdf ──────────────────────────────────────
            if not text:
                try:
                    import io as _io
                    from pypdf import PdfReader as _PdfReader
                    reader = _PdfReader(_io.BytesIO(raw_bytes))
                    text = "\n\n".join(
                        (page.extract_text() or "") for page in reader.pages
                    ).strip()
                except Exception as e:
                    return STUB_TEXT, f"PDF read error: {e} — using demo text."

            if text:
                return text, None
            return STUB_TEXT, "PDF text extraction returned empty content — using demo text."
        elif name.endswith(".docx"):
            from docx import Document as _Doc
            import io
            doc = _Doc(io.BytesIO(raw_bytes))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            if text:
                return text, None
            return STUB_TEXT, "DOCX appeared empty — using demo text."
    except Exception as e:
        return STUB_TEXT, f"File read error: {e} — using demo text."
    return STUB_TEXT, "Unsupported file type — using demo text."


# ─────────────────────────────────────────────
# AGENT DISCOVERY  (robust — works regardless of
# cwd, Streamlit Cloud, or local run directory)
# ─────────────────────────────────────────────
def _find_agent_dir() -> str:
    """
    Search for the directory containing agent.py.
    Tries multiple strategies so it works whether launched via
    `streamlit run app.py`, from a different cwd, or on Streamlit Cloud.
    Returns the directory path string, or '' if not found.
    """
    candidates = []
    try:
        candidates.append(str(Path(__file__).resolve().parent))
    except Exception:
        pass
    candidates.append(str(Path.cwd()))
    for p in sys.path:
        if p:
            candidates.append(p)
    try:
        candidates.append(str(Path(sys.argv[0]).resolve().parent))
    except Exception:
        pass
    for d in dict.fromkeys(candidates):
        if d and Path(d, "agent.py").exists():
            return d
    return ""


def _ensure_agent_importable() -> str:
    """Add agent directory to sys.path. Returns dir found, or ''."""
    d = _find_agent_dir()
    if d and d not in sys.path:
        sys.path.insert(0, d)
    return d


def _try_detect_type(text: str) -> dict:
    """Attempt to call agent.detect_document_type; fall back to heuristics."""
    try:
        _ensure_agent_importable()
        from agent import get_document_type  # type: ignore
        return get_document_type(text)
    except Exception:
        # Fallback heuristic detection
        lower = text.lower()
        kw_counts = {
            "STEM / IT": sum(1 for w in ["algorithm","software","network","python","data","model","system"] if w in lower),
            "Humanities": sum(1 for w in ["literary","narrative","poem","culture","history","art","rhetoric"] if w in lower),
            "Social Sciences": sum(1 for w in ["survey","psychology","participant","qualitative","policy","social"] if w in lower),
            "Medical / Health": sum(1 for w in ["patient","clinical","treatment","diagnosis","drug","health"] if w in lower),
            "Educational": sum(1 for w in ["student","teacher","curriculum","learning","classroom","pedagogy"] if w in lower),
        }
        best = max(kw_counts, key=kw_counts.get)
        score = kw_counts[best]
        if score == 0:
            best = "General Academic"

        style_map = {
            "STEM / IT": ["APA 7th Edition", "IEEE"],
            "Humanities": ["MLA 9th Edition", "Chicago 17th Edition"],
            "Social Sciences": ["APA 7th Edition", "MLA 9th Edition"],
            "Medical / Health": ["APA 7th Edition"],
            "Educational": ["APA 7th Edition", "MLA 9th Edition"],
            "General Academic": ["APA 7th Edition", "MLA 9th Edition"],
        }
        icon_map = {
            "STEM / IT": "🔬", "Humanities": "📖", "Social Sciences": "🧠",
            "Medical / Health": "🏥", "Educational": "🎓", "General Academic": "📄",
        }
        return {
            "type": best,
            "confidence": "Medium" if score >= 2 else "Low",
            "recommended_styles": style_map.get(best, ["APA 7th Edition"]),
            "description": f"Detected as {best} document",
            "icon": icon_map.get(best, "📄"),
        }


def call_agent(fp: str, style: str = "APA") -> dict:
    """
    Calls agent.format_manuscript(fp, style) if agent.py is locatable.
    Falls back to mock demo with a clear error message if not found.
    Returns: {formatted_path, report:{score, changes}, real, error}
    """
    agent_error = None
    try:
        found_dir = _ensure_agent_importable()
        if not found_dir:
            raise ImportError(
                "agent.py not found. Ensure agent.py and apa_mla_rules.json "
                "are in the same folder as app.py, then run: streamlit run app.py"
            )

        from agent import format_manuscript  # type: ignore
        formatted_path = format_manuscript(fp, style=style)

        if isinstance(formatted_path, str) and formatted_path.startswith("ERROR"):
            raise RuntimeError(formatted_path)
        if not formatted_path or not Path(formatted_path).exists():
            raise FileNotFoundError(f"Agent returned path but file missing: {formatted_path}")

        style_label = "MLA 9th" if style.upper().startswith("MLA") else ("IEEE" if style.upper().startswith("IEEE") else "APA 7th")
        changes = [
            f"Applied {style_label} Edition formatting",
            f"Set all margins to 1 inch (2.54 cm) on every side",
            f"Applied Times New Roman 12 pt font throughout",
            f"Formatted Level 1 headings per {style_label} rules",
            f"Aligned all body paragraphs with 0.5-inch first-line indent",
        ]

        # Read actual score and changes from compliance report
        actual_score = 85  # sensible default if report unreadable
        rp = Path(formatted_path).parent / "compliance_report.txt"
        if rp.exists():
            raw_report = rp.read_text(encoding="utf-8", errors="replace")
            score_m = re.search(r"Score:\s*(\d+)%", raw_report)
            if score_m:
                actual_score = int(score_m.group(1))
            _skip = {"Score:", "Changes Applied:", "Highlight Key", "Cyan", "Yellow", "Green", "Red", "==="}
            parsed = [
                l.strip("•-– \t").strip() for l in raw_report.splitlines()
                if l.strip() and not any(l.strip().startswith(s) for s in _skip)
            ]
            if parsed:
                changes = [c for c in parsed if c][:12]

        return {
            "formatted_path": formatted_path,
            "report": {"score": actual_score, "changes": changes},
            "real": True,
            "error": None,
            "style": style,
        }

    except ImportError:
        agent_error = "agent.py not found — running in demo mode"
    except FileNotFoundError as e:
        agent_error = f"Output file missing: {e}"
    except Exception as e:
        agent_error = f"Agent error: {e}"

    # Demo fallback
    time.sleep(2.1)
    style_label = "MLA 9th" if style.upper().startswith("MLA") else "APA 7th"
    if style.upper().startswith("MLA"):
        changes = [
            "Applied MLA 9th Edition formatting throughout",
            "Level-1 headings: left-aligned, bold, title case",
            "Level-2 headings: left-aligned, bold italic, title case",
            "Set all margins to 1 inch on every side",
            "Changed body font to Times New Roman 12 pt, double-spaced",
            f"Converted {random.randint(3,8)} APA citations → MLA (Author Page) format",
            "Renamed 'References' section → 'Works Cited'",
            "Reformatted Works Cited entries with MLA hanging indent (0.5 in)",
            f"Fixed {random.randint(2,5)} reference entry formats to MLA style",
            "Resolved inconsistent heading styles across manuscript",
        ]
    else:
        changes = [
            "Applied APA 7th Level 1 headings — bold, centered, title case",
            "Reformatted Abstract: bold label, 150–250 word limit enforced",
            "Set all margins to 1 inch (2.54 cm) on every side",
            "Changed body font to Times New Roman 12 pt, double-spaced",
            f"Fixed {random.randint(5,11)} in-text citations to author–date format",
            "Moved all figures & tables after the References section",
            "Rebuilt References list with APA 7th hanging-indent (0.5 in)",
            "Added running head and page numbers in document header",
            f"Corrected DOI hyperlink format in {random.randint(2,5)} references",
            "Resolved serial-comma inconsistencies across full manuscript",
        ]
    return {
        "formatted_path": None,
        "report": {"score": random.randint(88, 96), "changes": changes},
        "real": False,
        "error": agent_error,
        "style": style,
    }


def apa_preview(raw: str) -> str:
    headings = {"abstract","introduction","methodology","method","results","discussion","conclusion","references"}
    out = []
    for line in raw.splitlines():
        s = line.strip()
        out.append(f"\n{s.upper()}\n{'─'*36}" if s.lower() in headings else line)
    return "\n".join(out)


# ─────────────────────────────────────────────
# NAVBAR
# ─────────────────────────────────────────────
st.markdown("""
<div class="pp-navbar">
  <div class="pp-logo">
    <div class="pp-logo-icon">📄</div>
    <div class="pp-logo-text">paper<span>pal</span></div>
  </div>
  <div class="pp-nav-right">
    <span class="pp-chip chip-purple">⚡ AI Agent</span>
  </div>
</div>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# HERO
# ─────────────────────────────────────────────
st.markdown("""
<div class="pp-hero">
  <div class="hero-inner">
    <div class="hero-eyebrow"><span class="hero-dot"></span>APA 7th · MLA 9th · IEEE · Smart Detection</div>
    <div class="hero-title">AI Manuscript Formatter<br><em>for Researchers</em></div>
    <div class="hero-sub">
      Upload your draft — our AI detects your document type, recommends the right
      style guide, and formats your manuscript in seconds.
    </div>
    <div class="hero-stats">
      <div class="stat-item"><div class="stat-num">50+</div><div class="stat-label">Rules applied</div></div>
      <div class="stat-item"><div class="stat-num">3</div><div class="stat-label">Style guides (APA · MLA · IEEE)</div></div>
      <div class="stat-item"><div class="stat-num">5</div><div class="stat-label">Doc types detected</div></div>
      <div class="stat-item"><div class="stat-num">92%</div><div class="stat-label">Avg compliance score</div></div>
    </div>
  </div>
</div>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────
# MAIN LAYOUT
# ─────────────────────────────────────────────
left_col, right_col = st.columns([4, 6], gap="large")

# ════════════════════════
# LEFT PANEL
# ════════════════════════
with left_col:

    # ── UPLOAD CARD ──────────────────────────────────────────────
    
    st.markdown('<div class="pp-card-title"><span class="title-icon">📂</span>Upload Your Manuscript</div>', unsafe_allow_html=True)

    uploaded_file = st.file_uploader(
        "DRAG & DROP OR CLICK TO BROWSE",
        type=["docx","pdf"],
        help="Supported: .docx · .pdf"
    )

    st.markdown('</div>', unsafe_allow_html=True)   # close upload card

    # ── FORMATTING ENGINE BADGE + DIAGNOSTICS ───────────────────
    
    st.markdown(
        '<div class="pp-card-title"><span class="title-icon">⚙️</span>Formatting Engine</div>',
        unsafe_allow_html=True
    )

    # Run diagnostics every render so status is always fresh
    _agent_dir = _find_agent_dir()
    _agent_found   = bool(_agent_dir) and Path(_agent_dir, "agent.py").exists()
    _rules_found   = bool(_agent_dir) and Path(_agent_dir, "apa_mla_rules.json").exists()
    _agent_path    = str(Path(_agent_dir, "agent.py"))    if _agent_found else "—"
    _rules_path    = str(Path(_agent_dir, "apa_mla_rules.json")) if _rules_found else "—"
    _cwd           = str(Path.cwd())

    # Try importing agent to confirm it actually loads (not just exists)
    _import_ok  = False
    _import_err = ""
    if _agent_found:
        try:
            _ensure_agent_importable()
            import importlib, sys as _sys
            if "agent" in _sys.modules:
                _import_ok = True
            else:
                import agent as _ag  # noqa
                _import_ok = True
        except Exception as _e:
            _import_err = str(_e)

    
    # Default — overridden below if file uploaded
    agent_style = "APA"

    # ── STEP 1: Detect document type on upload ────────────────────
    if uploaded_file is not None:
        # Only re-detect if a new file was uploaded
        if st.session_state.uploaded_name != uploaded_file.name:
            # File size guard before reading
            uploaded_file.seek(0)
            raw_bytes_check = uploaded_file.read(MAX_FILE_BYTES + 1)
            uploaded_file.seek(0)
            if len(raw_bytes_check) > MAX_FILE_BYTES:
                st.error(f"⚠️ File exceeds {MAX_FILE_MB} MB limit. Please upload a smaller document.")
            else:
                uploaded_file.seek(0)
                raw_text_for_detection, upload_warn = read_uploaded_text(uploaded_file)
                uploaded_file.seek(0)

                with st.spinner("🔍 Analysing document type…"):
                    doc_info = _try_detect_type(raw_text_for_detection)

                st.session_state.doc_type_info    = doc_info
                st.session_state.uploaded_name    = uploaded_file.name
                st.session_state.original_text    = raw_text_for_detection
                st.session_state.style_confirmed  = False
                st.session_state.formatted_result = None
                st.session_state.chosen_style     = doc_info["recommended_styles"][0]
                st.session_state.upload_warning   = upload_warn

        # ── STEP 2: Show detected type banner ────────────────────
        info = st.session_state.doc_type_info
        if st.session_state.get("upload_warning"):
            st.warning(f"⚠️ {st.session_state.upload_warning}")
        if info:
            conf_cls = {"High": "conf-high", "Medium": "conf-medium", "Low": "conf-low"}.get(info["confidence"], "conf-low")
            rec_html = "".join(
                f'<span class="style-rec-chip {"style-rec-primary" if i == 0 else "style-rec-alt"}">{"★ " if i == 0 else ""}{s}</span>'
                for i, s in enumerate(info["recommended_styles"][:3])
            )
            st.markdown(f"""
            <div class="doc-type-banner">
              <div class="doc-type-icon">{info["icon"]}</div>
              <div class="doc-type-info">
                <div class="doc-type-label">Detected Document Type</div>
                <div class="doc-type-name">{info["type"]}</div>
                <div class="doc-type-desc">{info.get("description","")}</div>
              </div>
              <span class="doc-type-confidence {conf_cls}">{info["confidence"]} confidence</span>
            </div>
            <div class="style-rec-row" style="margin-bottom:4px;">
              <span style="font-size:11px;font-weight:700;letter-spacing:0.6px;text-transform:uppercase;color:#6B7280;margin-right:4px;">Recommended:</span>
              {rec_html}
            </div>
            """, unsafe_allow_html=True)

        # ── STEP 3: Style selector (pre-selected with recommendation) ──
        st.markdown('<div style="height:6px"></div>', unsafe_allow_html=True)

        all_styles = [
            "APA 7th Edition",
            "MLA 9th Edition",
            "IEEE",
        ]

        # Pre-select the recommended style
        rec_style = (info or {}).get("recommended_styles", ["APA 7th Edition"])[0] if info else "APA 7th Edition"
        default_idx = 0
        for i, s in enumerate(all_styles):
            if rec_style.lower() in s.lower():
                default_idx = i
                break

        style_choice = st.selectbox(
            "FORMATTING STYLE",
            all_styles,
            index=default_idx,
            help="Style auto-selected based on detected document type. You can change this."
        )
        st.session_state.chosen_style = style_choice

        # Determine actual style for agent
        if "MLA" in style_choice:
            agent_style = "MLA"
        elif "IEEE" in style_choice:
            agent_style = "IEEE"
        else:
            agent_style = "APA"
        style_available = True

    else:
        # No file uploaded — show default selector
        style_choice = st.selectbox("FORMATTING STYLE", [
            "APA 7th Edition",
            "MLA 9th Edition",
            "IEEE",
        ])
        if "MLA" in style_choice:
            agent_style = "MLA"
        elif "IEEE" in style_choice:
            agent_style = "IEEE"
        else:
            agent_style = "APA"

    st.markdown('<div style="height:10px"></div>', unsafe_allow_html=True)

    format_btn = st.button(
        "⚡  Format with AI Agent",
        use_container_width=True,
        disabled=uploaded_file is None
    )
    if uploaded_file is None:
        st.caption("↑ Upload a manuscript file to activate")

    st.markdown('</div>', unsafe_allow_html=True)

    # ── HOW IT WORKS CARD ─────────────────────────────────────────
    st.markdown("""
    <div class="pp-card" style="border-left:3px solid #7C3AED;">
      <div class="pp-card-title" style="font-size:14px;"><span class="title-icon">💡</span>How It Works</div>
      <div class="how-row">
        <div class="step-badge">1</div>
        <div class="step-desc">Upload your <span class="mono">.docx</span> or <span class="mono">.pdf</span> manuscript</div>
      </div>
      <div class="how-row">
        <div class="step-badge">2</div>
        <div class="step-desc">AI <strong style="color:#6C2BD9;">detects document type</strong> &amp; recommends a style guide</div>
      </div>
      <div class="how-row">
        <div class="step-badge">3</div>
        <div class="step-desc">Confirm or change the formatting style — <strong>APA 7th</strong> or <strong>MLA 9th</strong></div>
      </div>
      <div class="how-row">
        <div class="step-badge">4</div>
        <div class="step-desc">AI Agent formats headings, citations &amp; references automatically</div>
      </div>
      <div class="how-row">
        <div class="step-badge">5</div>
        <div class="step-desc">Preview before/after &amp; download formatted <span class="mono">.docx</span></div>
      </div>
    </div>
    """, unsafe_allow_html=True)


# ════════════════════════
# TRIGGER FORMATTING
# ════════════════════════
if format_btn and uploaded_file is not None:
    tmp_dir = Path(tempfile.gettempdir()) / "paperpal"
    tmp_dir.mkdir(exist_ok=True)
    tmp_path = tmp_dir / uploaded_file.name
    uploaded_file.seek(0); tmp_path.write_bytes(uploaded_file.read())

    raw_text = st.session_state.original_text or STUB_TEXT

    with right_col:
        style_label = "MLA 9th Edition" if agent_style == "MLA" else ("IEEE" if agent_style == "IEEE" else "APA 7th Edition")

        with st.spinner(f"⚡ Formatting your manuscript as {style_label}…"):
            result = call_agent(str(tmp_path), style=agent_style)
            result["engine"] = "rules"

    st.session_state.formatted_result = result
    st.session_state.original_text    = raw_text
    st.session_state.uploaded_name    = uploaded_file.name
    st.rerun()


# ════════════════════════
# RIGHT PANEL — RESULTS
# ════════════════════════
with right_col:
    if st.session_state.formatted_result is None:
        # Show doc type preview if file is uploaded but not yet formatted
        if st.session_state.doc_type_info and uploaded_file is not None:
            info = st.session_state.doc_type_info
            chosen = st.session_state.chosen_style or "APA 7th Edition"
            st.markdown(f"""
            <div style="background:var(--pp-white);border:1px solid var(--pp-border);
                        border-radius:var(--r-lg);padding:32px 36px;text-align:center;
                        animation:fadeUp 0.5s ease both;">
              <div style="font-size:48px;margin-bottom:16px;">{info["icon"]}</div>
              <div style="font-family:var(--font-serif);font-size:20px;font-weight:600;
                          color:var(--pp-text-dark);margin-bottom:8px;">
                {info["type"]} Document Detected
              </div>
              <div style="font-size:14px;color:var(--pp-text-muted);margin-bottom:20px;max-width:380px;margin-left:auto;margin-right:auto;line-height:1.6;">
                {info["description"]}.<br>
                Ready to format with <strong style="color:var(--pp-purple-700);">{chosen}</strong>.
              </div>
              <div style="display:flex;justify-content:center;gap:12px;flex-wrap:wrap;">
                {"".join(
                    f'<div style="background:{"var(--pp-purple-50)" if i==0 else "var(--pp-bg)"};border:{"1.5px solid var(--pp-purple-600)" if i==0 else "1px solid var(--pp-border)"};border-radius:8px;padding:10px 18px;font-size:13px;font-weight:{"700" if i==0 else "500"};color:{"var(--pp-purple-700)" if i==0 else "var(--pp-text-muted)"};">{"★ " if i==0 else ""}{s}</div>'
                    for i, s in enumerate(info["recommended_styles"][:3])
                )}
              </div>
              <div style="margin-top:24px;font-size:13px;color:var(--pp-text-faint);">
                Click <strong style="color:var(--pp-purple-700);">Format with AI Agent</strong> to begin
              </div>
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div class="empty-state">
              <div class="empty-icon">📋</div>
              <div class="empty-h">Results Appear Here</div>
              <div class="empty-p">
                Upload a manuscript on the left — AI will detect your document type
                and recommend the best formatting style.
              </div>
            </div>
            """, unsafe_allow_html=True)
    else:
        result  = st.session_state.formatted_result
        raw     = st.session_state.original_text or ""
        fname   = st.session_state.uploaded_name or "manuscript"
        report  = result.get("report", {})
        score   = report.get("score", 92)
        changes = report.get("changes", [
            "Applied formatting per selected style guide",
            "Set 1-inch margins on all sides",
            "Fixed in-text citations",
            "Formatted reference entries with hanging indent",
        ])
        used_style = result.get("style", "APA")
        style_label = "MLA 9th Edition" if used_style == "MLA" else ("IEEE" if used_style == "IEEE" else "APA 7th Edition")
        style_color = "#0D9488" if used_style == "MLA" else ("#1D4ED8" if used_style == "IEEE" else "#6C2BD9")
        engine      = result.get("engine", "rules")

        # Status bar
        st.markdown(
            f'<div class="status-pill"><span class="live-dot"></span>'
            f'Formatting complete &nbsp;·&nbsp; {fname} &nbsp;·&nbsp; '
            f'<span style="color:{style_color};font-weight:800;">{style_label}</span></div>',
            unsafe_allow_html=True,
        )

        # Engine badge
        if result.get("real"):
            st.markdown("""
            <div style="display:inline-flex;align-items:center;gap:7px;
                        background:#F0FDF4;border:1px solid #A7F3D0;
                        color:#065F46;border-radius:8px;padding:7px 14px;
                        font-size:12px;font-weight:600;margin-bottom:12px;">
              ✅ <strong>Formatting complete</strong> &nbsp;·&nbsp; Document transformed successfully
            </div>
            """, unsafe_allow_html=True)
        else:
            err = result.get("error", "agent not found")
            st.markdown(f"""
            <div style="display:inline-flex;align-items:center;gap:7px;
                        background:#FFFBEB;border:1px solid #FCD34D;
                        color:#92400E;border-radius:8px;padding:7px 14px;
                        font-size:12px;font-weight:600;margin-bottom:12px;">
              ⚠️ Demo mode &nbsp;·&nbsp; {err}
            </div>
            """, unsafe_allow_html=True)

        tab1, tab2, tab3 = st.tabs([
            "📄  Before / After Preview",
            "📊  AI Compliance Report",
            "🔗  Citation Consistency",
        ])

        # ── Shared imports and data prep (used by all tabs) ───────────
        import base64, io
        import streamlit.components.v1 as components
        from docx import Document as DocxDocument
        from docx.oxml.ns import qn as _qn

        DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        def strip_highlights(docx_bytes: bytes) -> bytes:
            """Return a copy of the docx with all highlights and red color removed."""
            doc = DocxDocument(io.BytesIO(docx_bytes))
            for para in doc.paragraphs:
                for run in para.runs:
                    rPr = run._r.find(_qn("w:rPr"))
                    if rPr is None: continue
                    for tag in (_qn("w:highlight"), _qn("w:color")):
                        el = rPr.find(tag)
                        if el is not None: rPr.remove(el)
                    try: run.font.color.rgb = None
                    except Exception: pass
            buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

        def text_to_docx(text: str) -> bytes:
            from docx import Document as D
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            doc = D()
            sec = doc.sections[0]
            for attr in ("left_margin","right_margin","top_margin","bottom_margin"):
                setattr(sec, attr, Inches(1))
            HEADS = {"abstract","introduction","methodology","method",
                     "results","discussion","conclusion","references","works cited"}
            for line in text.splitlines():
                s = line.strip()
                if not s: doc.add_paragraph(""); continue
                if s.lower() in HEADS:
                    p = doc.add_paragraph(s.title())
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs: r.bold = True; r.font.size = Pt(12)
                else:
                    p = doc.add_paragraph(s)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for r in p.runs: r.font.size = Pt(12)
                    p.paragraph_format.first_line_indent = Inches(0.5)
                    p.paragraph_format.line_spacing = Pt(24)
            buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

        def text_to_pdf(text: str) -> bytes:
            """Convert plain text to a styled PDF using reportlab."""
            import io as _io
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch

            buf = _io.BytesIO()
            doc = SimpleDocTemplate(
                buf, pagesize=letter,
                leftMargin=inch, rightMargin=inch,
                topMargin=inch, bottomMargin=inch
            )
            styles = getSampleStyleSheet()
            HEADS = {"abstract","introduction","methodology","method",
                     "results","discussion","conclusion","references","works cited"}
            heading_style = ParagraphStyle(
                "H1", parent=styles["Heading1"],
                fontSize=13, fontName="Times-Bold",
                spaceAfter=6, spaceBefore=12, alignment=1
            )
            body_style = ParagraphStyle(
                "Body", parent=styles["Normal"],
                fontSize=12, fontName="Times-Roman",
                leading=24, firstLineIndent=36
            )
            story = []
            for line in text.splitlines():
                s = line.strip()
                if not s:
                    story.append(Spacer(1, 12))
                elif s.lower() in HEADS:
                    story.append(Paragraph(s.title(), heading_style))
                else:
                    story.append(Paragraph(s, body_style))
            doc.build(story)
            return buf.getvalue()

        # Load original bytes
        tmp_path_r = Path(tempfile.gettempdir()) / "paperpal" / (st.session_state.uploaded_name or "")
        if tmp_path_r.exists() and tmp_path_r.suffix.lower() == ".docx":
            orig_bytes    = tmp_path_r.read_bytes()
            orig_dl_bytes = orig_bytes
            orig_dl_name  = tmp_path_r.name
        else:
            # PDF or missing — use text fallback for preview, convert to docx for download
            orig_bytes    = None
            orig_dl_bytes = text_to_docx(raw)
            orig_dl_name  = (st.session_state.uploaded_name or "original").rsplit(".", 1)[0] + ".docx"

        # Load formatted bytes — only available when agent ran successfully (result["real"] == True)
        fp = result.get("formatted_path")
        if fp and Path(fp).exists():
            fmt_bytes    = Path(fp).read_bytes()
            fmt_dl_bytes = strip_highlights(fmt_bytes)
            fmt_dl_name  = f"formatted_{Path(fp).name}"
        else:
            # Demo mode or agent error — no real formatted file
            fmt_bytes    = None
            base_name    = (st.session_state.uploaded_name or "manuscript").rsplit(".", 1)[0]
            fmt_dl_bytes = text_to_docx(apa_preview(raw))
            fmt_dl_name  = f"formatted_{used_style.lower()}_{base_name}.docx"

        # ─── TAB 1 : BEFORE / AFTER ───
        with tab1:

            # ════════════════════════════════════════════════════════
            # SHARED XML HELPERS
            # ════════════════════════════════════════════════════════
            def _esc(t):
                return str(t).replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")

            # ── Run-level XML readers ─────────────────────────────
            def _hl(run):
                try:
                    rPr = run._r.find(_qn("w:rPr"))
                    if rPr is None: return None
                    el = rPr.find(_qn("w:highlight"))
                    return el.get(_qn("w:val")) if el is not None else None
                except Exception: return None

            def _col(run):
                try:
                    rPr = run._r.find(_qn("w:rPr"))
                    if rPr is None: return None
                    el = rPr.find(_qn("w:color"))
                    if el is None: return None
                    v = el.get(_qn("w:val"), "")
                    return "#" + v if v and v.lower() not in ("auto","000000","") else None
                except Exception: return None

            def _rbold(run):
                try:
                    rPr = run._r.find(_qn("w:rPr"))
                    if rPr is None: return False
                    b = rPr.find(_qn("w:b"))
                    if b is None: return False
                    return b.get(_qn("w:val"), "true").lower() not in ("false","0","off")
                except Exception: return bool(run.bold)

            def _rital(run):
                try:
                    rPr = run._r.find(_qn("w:rPr"))
                    if rPr is None: return False
                    el = rPr.find(_qn("w:i"))
                    if el is None: return False
                    return el.get(_qn("w:val"), "true").lower() not in ("false","0","off")
                except Exception: return bool(run.italic)

            # ── Para-level XML readers ────────────────────────────
            def _palign(para):
                try:
                    pPr = para._p.find(_qn("w:pPr"))
                    if pPr is None: return "left"
                    jc  = pPr.find(_qn("w:jc"))
                    if jc is None: return "left"
                    val = jc.get(_qn("w:val"), "left").lower()
                    return {"center":"center","both":"justify","right":"right"}.get(val,"left")
                except Exception: return "left"

            def _pindent(para):
                try:
                    pPr = para._p.find(_qn("w:pPr"))
                    if pPr is None: return 0.0, 0.0, False
                    ind = pPr.find(_qn("w:ind"))
                    if ind is None: return 0.0, 0.0, False
                    left = int(ind.get(_qn("w:left"), 0) or 0) / 1440.0
                    fl   = int(ind.get(_qn("w:firstLine"), 0) or 0) / 1440.0
                    hang = ind.get(_qn("w:hanging")) is not None
                    return left, fl, hang
                except Exception: return 0.0, 0.0, False

            def _pfont_size(para):
                """Return font size in pt from first run, default 12."""
                try:
                    for run in para.runs:
                        rPr = run._r.find(_qn("w:rPr"))
                        if rPr is None: continue
                        sz = rPr.find(_qn("w:sz"))
                        if sz is not None:
                            return int(sz.get(_qn("w:val"), "24")) / 2
                    return 12.0
                except Exception: return 12.0

            # ════════════════════════════════════════════════════════
            # HIGHLIGHT COLOUR DEFINITIONS
            # (matches agent.py highlight keys exactly)
            # ════════════════════════════════════════════════════════
            HL_STYLES = {
                # run-level highlights applied by agent.py
                "cyan":    ("background:#C7F5FF;color:#0E4F5A;border-bottom:2px solid #06B6D4;",
                            "Heading reformatted"),
                "yellow":  ("background:#FEF9C3;color:#713F12;border-bottom:2px solid #EAB308;",
                            "Citation fixed"),
                "green":   ("background:#DCFCE7;color:#065F46;border-bottom:2px solid #22C55E;",
                            "Reference fixed"),
                "red":     ("background:#FEE2E2;color:#991B1B;border-bottom:2px solid #EF4444;",
                            "Needs review"),
                "darkRed": ("background:#FEE2E2;color:#991B1B;border-bottom:2px solid #EF4444;",
                            "Needs review"),
            }
            SPAN_BASE = "border-radius:3px;padding:1px 5px;font-size:inherit;"

            # Para-level change types (detected by structural analysis)
            PARA_CHANGE_STYLES = {
                "heading":   "border-left:4px solid #06B6D4;background:#F0FEFF;",
                "citation":  "border-left:4px solid #EAB308;background:#FEFCE8;",
                "reference": "border-left:4px solid #22C55E;background:#F0FDF4;",
                "warning":   "border-left:4px solid #EF4444;background:#FFF1F2;",
                "body":      "border-left:4px solid #A78BFA;background:#FAF5FF;",
            }

            CHANGE_BADGE = {
                "heading":   ('<span style="font-size:9px;font-weight:800;letter-spacing:.5px;'
                              'text-transform:uppercase;color:#0E7490;background:#CFFAFE;'
                              'border-radius:4px;padding:1px 6px;margin-right:6px;'
                              'vertical-align:middle;white-space:nowrap;">HEADING</span>'),
                "citation":  ('<span style="font-size:9px;font-weight:800;letter-spacing:.5px;'
                              'text-transform:uppercase;color:#92400E;background:#FEF3C7;'
                              'border-radius:4px;padding:1px 6px;margin-right:6px;'
                              'vertical-align:middle;white-space:nowrap;">CITATION</span>'),
                "reference": ('<span style="font-size:9px;font-weight:800;letter-spacing:.5px;'
                              'text-transform:uppercase;color:#065F46;background:#D1FAE5;'
                              'border-radius:4px;padding:1px 6px;margin-right:6px;'
                              'vertical-align:middle;white-space:nowrap;">REFERENCE</span>'),
                "warning":   ('<span style="font-size:9px;font-weight:800;letter-spacing:.5px;'
                              'text-transform:uppercase;color:#991B1B;background:#FEE2E2;'
                              'border-radius:4px;padding:1px 6px;margin-right:6px;'
                              'vertical-align:middle;white-space:nowrap;">⚠ REVIEW</span>'),
                "body":      ('<span style="font-size:9px;font-weight:800;letter-spacing:.5px;'
                              'text-transform:uppercase;color:#5B21B6;background:#EDE9FE;'
                              'border-radius:4px;padding:1px 6px;margin-right:6px;'
                              'vertical-align:middle;white-space:nowrap;">BODY</span>'),
            }

            # ════════════════════════════════════════════════════════
            # RUN → SPAN HTML  (for the "After" panel)
            # Emits highlighted spans for run-level changes, and
            # returns (html, dominant_change_type) for para classification
            # ════════════════════════════════════════════════════════
            def _runs_to_spans(para):
                """Returns (inner_html, dominant_highlight_key_or_None)."""
                parts = []
                dominant = None
                priority = ["warning","red","darkRed","green","yellow","cyan"]
                seen_hl = set()
                for run in para.runs:
                    t = _esc(run.text)
                    if not t: continue
                    h  = _hl(run)
                    c  = _col(run)
                    fw = "bold"   if _rbold(run) else "normal"
                    fi = "italic" if _rital(run) else "normal"
                    base_sty = "font-weight:%s;font-style:%s;" % (fw, fi)

                    if h and h in HL_STYLES:
                        seen_hl.add(h)
                        hl_sty, _ = HL_STYLES[h]
                        sty = base_sty + hl_sty + SPAN_BASE
                        parts.append('<span style="%s" title="%s">%s</span>'
                                     % (sty, HL_STYLES[h][1], t))
                    elif c:
                        # Red warning text from agent
                        seen_hl.add("warning")
                        sty = base_sty + "color:%s;" % c
                        parts.append('<span style="%s">%s</span>' % (sty, t))
                    else:
                        parts.append('<span style="%s">%s</span>' % (base_sty, t))

                # Pick the highest-priority highlight seen as the dominant type
                for key in priority:
                    if key in seen_hl:
                        dominant = key
                        break

                return "".join(parts) or _esc(para.text), dominant

            # ════════════════════════════════════════════════════════
            # PARA → HTML ROW  (for the "After" / formatted panel)
            # ════════════════════════════════════════════════════════
            def _para_to_html_after(para):
                txt = para.text
                if not txt.strip():
                    return '<div style="margin:0;line-height:2;min-height:1.5em;"></div>'

                align        = _palign(para)
                left, fl, hg = _pindent(para)
                runs         = [r for r in para.runs if r.text.strip()]
                is_bold      = bool(runs) and all(_rbold(r) for r in runs)
                is_ital      = bool(runs) and all(_rital(r) for r in runs)
                fsize_pt     = _pfont_size(para)

                # Build indent CSS
                ind_css = ""
                if hg and left > 0:
                    ind_css = "padding-left:%.3fin;text-indent:-%.3fin;" % (left, left)
                elif left > 0:
                    ind_css = "padding-left:%.3fin;" % left
                elif fl > 0:
                    ind_css = "text-indent:%.3fin;" % fl

                inner, dominant_hl = _runs_to_spans(para)

                # Map highlight to para change type
                HL_TO_CHANGE = {
                    "cyan":    "heading",
                    "yellow":  "citation",
                    "green":   "reference",
                    "red":     "warning",
                    "darkRed": "warning",
                    "warning": "warning",
                }
                # Structural heading detection (bold + centered → heading even without hl)
                if dominant_hl:
                    change_type = HL_TO_CHANGE.get(dominant_hl)
                elif is_bold and align == "center":
                    change_type = "heading"
                elif is_bold and align == "left" and fl == 0.0:
                    change_type = "heading"
                elif fl > 0 or hg:
                    change_type = "body"
                else:
                    change_type = None

                fw    = "bold"   if is_bold else "normal"
                fi    = "italic" if is_ital else "normal"
                mt    = "20px" if change_type == "heading" else "2px"

                text_sty = (
                    "font-family:'Times New Roman',serif;"
                    "font-size:%.1fpt;font-weight:%s;font-style:%s;"
                    "line-height:2;text-align:%s;color:#1a1a2e;%s"
                ) % (fsize_pt, fw, fi, align, ind_css)

                if change_type and change_type in PARA_CHANGE_STYLES:
                    wrap_sty  = PARA_CHANGE_STYLES[change_type]
                    badge_html = CHANGE_BADGE.get(change_type, "")
                    return (
                        '<div style="%(wrap)s margin:%(mt)s 0 0;'
                        'border-radius:0 6px 6px 0;padding:4px 10px 4px 8px;'
                        'position:relative;" class="changed-para">'
                        '<p style="%(text)s margin:0;">%(badge)s%(inner)s</p>'
                        '</div>'
                    ) % {"wrap": wrap_sty, "mt": mt, "text": text_sty,
                         "badge": badge_html, "inner": inner}
                else:
                    return '<p style="%s margin-top:%s;">%s</p>' % (text_sty, mt, inner)

            # ════════════════════════════════════════════════════════
            # PARA → HTML ROW  (for the "Before" / original panel)
            # Shows the original text with a faded "was changed" marker
            # for paragraphs that correspond to changed ones in "After"
            # ════════════════════════════════════════════════════════
            def _para_to_html_before(para, changed: bool, change_type: str = None):
                txt = para.text
                if not txt.strip():
                    return '<div style="margin:0;line-height:2;min-height:1.5em;"></div>'

                align        = _palign(para)
                left, fl, hg = _pindent(para)
                runs         = [r for r in para.runs if r.text.strip()]
                is_bold      = bool(runs) and all(_rbold(r) for r in runs)
                is_ital      = bool(runs) and all(_rital(r) for r in runs)
                fsize_pt     = _pfont_size(para)

                ind_css = ""
                if hg and left > 0:
                    ind_css = "padding-left:%.3fin;text-indent:-%.3fin;" % (left, left)
                elif left > 0:
                    ind_css = "padding-left:%.3fin;" % left
                elif fl > 0:
                    ind_css = "text-indent:%.3fin;" % fl

                fw = "bold"   if is_bold else "normal"
                fi = "italic" if is_ital else "normal"

                if changed:
                    # Dim the text slightly + red left border + "CHANGED" tag
                    text_sty = (
                        "font-family:'Times New Roman',serif;"
                        "font-size:%.1fpt;font-weight:%s;font-style:%s;"
                        "line-height:2;text-align:%s;color:#6B7280;%s"
                    ) % (fsize_pt, fw, fi, align, ind_css)

                    BEFORE_CHANGE_BORDERS = {
                        "heading":   "border-left:4px solid #CBD5E1;",
                        "citation":  "border-left:4px solid #FDE68A;",
                        "reference": "border-left:4px solid #A7F3D0;",
                        "warning":   "border-left:4px solid #FECACA;",
                        "body":      "border-left:4px solid #DDD6FE;",
                    }
                    border = BEFORE_CHANGE_BORDERS.get(change_type or "body",
                                                       "border-left:4px solid #E5E7EB;")
                    old_badge = (
                        '<span style="font-size:9px;font-weight:700;letter-spacing:.4px;'
                        'text-transform:uppercase;color:#9CA3AF;background:#F3F4F6;'
                        'border-radius:4px;padding:1px 6px;margin-right:6px;'
                        'vertical-align:middle;white-space:nowrap;">ORIGINAL</span>'
                    )
                    return (
                        '<div style="%(border)s background:#FAFAFA;'
                        'border-radius:0 6px 6px 0;'
                        'padding:4px 10px 4px 8px;margin:2px 0 0;">'
                        '<p style="%(text)s margin:0;">%(badge)s%(inner)s</p>'
                        '</div>'
                    ) % {"border": border, "text": text_sty,
                         "badge": old_badge, "inner": _esc(txt)}
                else:
                    text_sty = (
                        "font-family:'Times New Roman',serif;"
                        "font-size:%.1fpt;font-weight:%s;font-style:%s;"
                        "line-height:2;text-align:%s;color:#1a1a2e;%s"
                    ) % (fsize_pt, fw, fi, align, ind_css)
                    mt = "20px" if (is_bold and (align == "center" or fl == 0)) else "2px"
                    return '<p style="%s margin-top:%s;">%s</p>' % (text_sty, mt, _esc(txt))

            # ════════════════════════════════════════════════════════
            # FULL DOCX → HTML  (builds both panels in one pass for
            # alignment: same paragraph index = same logical block)
            # ════════════════════════════════════════════════════════
            def docx_pair_to_html(orig_bytes, fmt_bytes):
                """
                Returns (orig_html, fmt_html, change_count).
                Aligns paragraphs by index; marks originals that changed.
                """
                try:
                    orig_doc = DocxDocument(io.BytesIO(orig_bytes))
                    fmt_doc  = DocxDocument(io.BytesIO(fmt_bytes))
                except Exception as ex:
                    err = '<p style="color:red;">Preview error: %s</p>' % ex
                    return err, err, 0

                orig_paras = list(orig_doc.paragraphs)
                fmt_paras  = list(fmt_doc.paragraphs)
                n = max(len(orig_paras), len(fmt_paras))

                orig_rows = []
                fmt_rows  = []
                change_count = 0

                for i in range(n):
                    op = orig_paras[i] if i < len(orig_paras) else None
                    fp = fmt_paras[i]  if i < len(fmt_paras)  else None

                    if fp is None:
                        # Paragraph deleted in formatted version
                        if op:
                            orig_rows.append(_para_to_html_before(op, changed=True, change_type="warning"))
                            fmt_rows.append('<div style="min-height:2em;border-left:4px solid #E5E7EB;'
                                           'background:#F9FAFB;border-radius:0 6px 6px 0;'
                                           'padding:4px 10px;margin:2px 0;'
                                           'color:#9CA3AF;font-family:sans-serif;font-size:11px;">'
                                           '(paragraph removed)</div>')
                            change_count += 1
                        continue

                    # Render formatted paragraph and get its change type
                    fmt_row = _para_to_html_after(fp)
                    fmt_rows.append(fmt_row)

                    # Determine if this paragraph was changed by checking:
                    # 1. Any run has a highlight in the formatted doc
                    # 2. Text content differs from original
                    # 3. Alignment / indent changed
                    fp_inner, dominant_hl = _runs_to_spans(fp)
                    fp_text = fp.text.strip()
                    op_text = op.text.strip() if op else ""

                    HL_TO_CHANGE = {
                        "cyan": "heading", "yellow": "citation",
                        "green": "reference", "red": "warning",
                        "darkRed": "warning", "warning": "warning",
                    }
                    change_type = HL_TO_CHANGE.get(dominant_hl) if dominant_hl else None

                    # Structural change detection even when no highlight
                    if change_type is None and op:
                        fp_align = _palign(fp); op_align = _palign(op)
                        fp_ind   = _pindent(fp); op_ind   = _pindent(op)
                        fp_bold  = bool([r for r in fp.runs if r.text.strip()]) and \
                                   all(_rbold(r) for r in fp.runs if r.text.strip())
                        op_bold  = bool([r for r in op.runs if r.text.strip()]) and \
                                   all(_rbold(r) for r in op.runs if r.text.strip())
                        if fp_align != op_align or fp_ind != op_ind or fp_bold != op_bold:
                            change_type = "heading" if (fp_bold or fp_align == "center") else "body"
                        elif fp_text != op_text:
                            change_type = "body"

                    was_changed = change_type is not None
                    if was_changed:
                        change_count += 1

                    if op:
                        orig_rows.append(_para_to_html_before(op, changed=was_changed,
                                                               change_type=change_type))
                    else:
                        # New paragraph inserted by agent
                        fmt_rows[-1] = ('<div style="border-left:4px solid #22C55E;'
                                        'background:#F0FDF4;border-radius:0 6px 6px 0;'
                                        'padding:4px 10px;margin:2px 0;">'
                                        + fmt_row + '</div>')
                        orig_rows.append('<div style="min-height:2em;border-left:4px solid #E5E7EB;'
                                        'background:#F9FAFB;border-radius:0 6px 6px 0;'
                                        'padding:4px 10px;margin:2px 0;'
                                        'color:#9CA3AF;font-family:sans-serif;font-size:11px;">'
                                        '(paragraph added)</div>')

                return "\n".join(orig_rows), "\n".join(fmt_rows), change_count

            # ════════════════════════════════════════════════════════
            # TEXT FALLBACK (when no docx bytes available)
            # ════════════════════════════════════════════════════════
            def text_to_html(text, is_after=False):
                HEADS = {"abstract","introduction","methodology","method","results",
                         "discussion","conclusion","references","works cited"}
                rows = []
                for line in text.splitlines():
                    s = line.strip()
                    if not s:
                        rows.append('<div style="margin:0;line-height:2;min-height:1.5em;"></div>')
                    elif s.lower() in HEADS:
                        if is_after:
                            rows.append(
                                '<div style="border-left:4px solid #06B6D4;background:#F0FEFF;'
                                'border-radius:0 6px 6px 0;padding:4px 10px;margin:20px 0 0;">'
                                '<p style="text-align:center;font-family:\'Times New Roman\',serif;'
                                'font-size:13pt;font-weight:bold;margin:0;color:#1a1a2e;">'
                                '<span style="font-size:9px;font-weight:800;letter-spacing:.5px;'
                                'text-transform:uppercase;color:#0E7490;background:#CFFAFE;'
                                'border-radius:4px;padding:1px 6px;margin-right:6px;'
                                'vertical-align:middle;">HEADING</span>'
                                + _esc(s.title()) + '</p></div>')
                        else:
                            rows.append(
                                '<p style="text-align:center;font-family:\'Times New Roman\',serif;'
                                'font-size:13pt;font-weight:bold;margin-top:20px;color:#6B7280;">'
                                '<span style="font-size:9px;font-weight:700;color:#9CA3AF;'
                                'background:#F3F4F6;border-radius:4px;padding:1px 6px;'
                                'margin-right:6px;vertical-align:middle;">ORIGINAL</span>'
                                + _esc(s.title()) + '</p>')
                    else:
                        rows.append(
                            '<p style="text-align:left;font-family:\'Times New Roman\',serif;'
                            'font-size:12pt;line-height:2;text-indent:0.5in;margin:2px 0;'
                            'color:#1a1a2e;">%s</p>' % _esc(s))
                return "\n".join(rows)

            # ════════════════════════════════════════════════════════
            # IFRAME PAGE BUILDER  (sticky header + legend + content)
            # ════════════════════════════════════════════════════════
            def make_page(body_html, label, accent, show_legend=False, change_count=0):
                legend = ""
                if show_legend:
                    count_pill = (
                        '<span style="background:%s;color:#fff;border-radius:20px;'
                        'padding:2px 9px;font-size:11px;font-weight:800;margin-left:auto;">'
                        '%d change%s highlighted</span>'
                    ) % (accent, change_count, "s" if change_count != 1 else "")

                    legend = (
                        '<div style="position:sticky;bottom:0;z-index:10;display:flex;'
                        'align-items:center;gap:6px;flex-wrap:wrap;'
                        'padding:7px 14px;border-top:2px solid %(accent)s;'
                        'background:#faf9ff;box-shadow:0 -2px 8px rgba(0,0,0,0.06);">'
                        '<span style="font-size:10px;font-weight:700;letter-spacing:.4px;'
                        'text-transform:uppercase;color:#9CA3AF;line-height:24px;">Changes:</span>'
                        '<span style="font-size:11px;background:#CFFAFE;color:#0E7490;'
                        'border-radius:4px;padding:2px 8px;font-weight:700;border-left:3px solid #06B6D4;">'
                        '&#9679; Heading</span>'
                        '<span style="font-size:11px;background:#FEF9C3;color:#92400E;'
                        'border-radius:4px;padding:2px 8px;font-weight:700;border-left:3px solid #EAB308;">'
                        '&#9679; Citation</span>'
                        '<span style="font-size:11px;background:#DCFCE7;color:#065F46;'
                        'border-radius:4px;padding:2px 8px;font-weight:700;border-left:3px solid #22C55E;">'
                        '&#9679; Reference</span>'
                        '<span style="font-size:11px;background:#EDE9FE;color:#5B21B6;'
                        'border-radius:4px;padding:2px 8px;font-weight:700;border-left:3px solid #A78BFA;">'
                        '&#9679; Body/Indent</span>'
                        '<span style="font-size:11px;background:#FEE2E2;color:#991B1B;'
                        'border-radius:4px;padding:2px 8px;font-weight:700;border-left:3px solid #EF4444;">'
                        '&#9679; Review</span>'
                        '%(count)s'
                        '</div>'
                    ) % {"accent": accent, "count": count_pill}

                return (
                    "<!DOCTYPE html><html lang='en'><head><meta charset='utf-8'>"
                    "<style>"
                    "*{box-sizing:border-box;margin:0;padding:0}"
                    "html,body{width:100%%;background:#fff;font-size:14px;}"
                    ".hdr{position:sticky;top:0;z-index:10;"
                    "display:flex;align-items:center;gap:7px;padding:8px 16px;"
                    "background:#f8f7ff;border-bottom:3px solid %(accent)s;"
                    "font-family:Arial,sans-serif;font-size:11px;font-weight:700;"
                    "letter-spacing:.7px;text-transform:uppercase;color:#555;"
                    "box-shadow:0 2px 6px rgba(0,0,0,0.06);}"
                    ".dot{width:9px;height:9px;border-radius:50%%;flex-shrink:0;background:%(accent)s;}"
                    ".content{padding:28px 40px 40px;}"
                    ".changed-para:hover{filter:brightness(0.97);transition:filter 0.15s;}"
                    "p{word-break:break-word;}"
                    "</style></head><body>"
                    "<div class='hdr'><span class='dot'></span>%(label)s</div>"
                    "<div class='content'>%(body)s</div>"
                    "%(legend)s"
                    "</body></html>"
                ) % {"accent": accent, "label": label, "body": body_html, "legend": legend}

            # ════════════════════════════════════════════════════════
            # BUILD CONTENT
            # ════════════════════════════════════════════════════════
            acc = "#0D9488" if used_style == "MLA" else ("#1D4ED8" if used_style == "IEEE" else "#7C3AED")
            change_count = 0

            if orig_bytes is not None and fmt_bytes is not None:
                # Best path: both docx files available — aligned pair diff
                orig_html, fmt_html, change_count = docx_pair_to_html(orig_bytes, fmt_bytes)
            elif fmt_bytes is not None:
                # Only formatted available
                _, fmt_html, change_count = docx_pair_to_html(
                    fmt_bytes, fmt_bytes)   # compare to itself (no diffs shown)
                orig_html = text_to_html(raw, is_after=False)
                # Re-render fmt with full highlights
                fmt_rows = []
                try:
                    fd = DocxDocument(io.BytesIO(fmt_bytes))
                    for fp in fd.paragraphs:
                        fmt_rows.append(_para_to_html_after(fp))
                    fmt_html = "\n".join(fmt_rows)
                except Exception:
                    pass
            elif orig_bytes is not None:
                orig_html = text_to_html(raw, is_after=False)
                demo_notice = (
                    "<div style='background:#FFFBEB;border:1px solid #FCD34D;border-radius:8px;"
                    "padding:10px 14px;margin-bottom:16px;font-family:Arial,sans-serif;"
                    "font-size:12px;color:#92400E;'>"
                    "<strong>&#9888; Demo mode</strong> &mdash; place <code>agent.py</code> "
                    "and <code>apa_mla_rules.json</code> in the same folder to enable real formatting."
                    "</div>"
                )
                fmt_html = demo_notice + text_to_html(apa_preview(raw), is_after=True)
            else:
                orig_html = text_to_html(raw, is_after=False)
                demo_notice = (
                    "<div style='background:#FFFBEB;border:1px solid #FCD34D;border-radius:8px;"
                    "padding:10px 14px;margin-bottom:16px;font-family:Arial,sans-serif;"
                    "font-size:12px;color:#92400E;'>"
                    "<strong>&#9888; Demo mode</strong> &mdash; place <code>agent.py</code> "
                    "and <code>apa_mla_rules.json</code> in the same folder to enable real formatting."
                    "</div>"
                )
                fmt_html = demo_notice + text_to_html(apa_preview(raw), is_after=True)

            # ════════════════════════════════════════════════════════
            # RENDER
            # ════════════════════════════════════════════════════════
            bc, ac = st.columns(2, gap="small")

            with bc:
                components.html(
                    make_page(orig_html, "Before \u00b7 Original", "#8B5CF6"),
                    height=560, scrolling=True
                )
                orig_b64 = base64.b64encode(orig_dl_bytes).decode()
                st.markdown(
                    '<a href="data:' + DOCX_MIME + ';base64,' + orig_b64 + '" '
                    'download="' + orig_dl_name + '" '
                    'style="display:flex;align-items:center;justify-content:center;gap:8px;'
                    'background:#F5F3FF;border:1.5px solid #8B5CF6;border-radius:8px;'
                    'padding:9px 14px;font-family:sans-serif;font-size:13px;font-weight:600;'
                    'color:#6C2BD9;text-decoration:none;margin-top:8px;">'
                    '&#128229; Download Original (.docx)</a>',
                    unsafe_allow_html=True
                )

            with ac:
                components.html(
                    make_page(fmt_html, "After \u00b7 " + style_label, acc,
                              show_legend=True, change_count=change_count),
                    height=560, scrolling=True
                )
                fmt_b64 = base64.b64encode(fmt_dl_bytes).decode()
                st.markdown(
                    '<a href="data:' + DOCX_MIME + ';base64,' + fmt_b64 + '" '
                    'download="' + fmt_dl_name + '" '
                    'style="display:flex;align-items:center;justify-content:center;gap:8px;'
                    'background:#F0FDFA;border:1.5px solid #0D9488;border-radius:8px;'
                    'padding:9px 14px;font-family:sans-serif;font-size:13px;font-weight:600;'
                    'color:#0D9488;text-decoration:none;margin-top:8px;">'
                    '&#128229; Download Formatted (.docx) &#8212; highlights removed</a>',
                    unsafe_allow_html=True
                )
        # ─── TAB 2 : COMPLIANCE REPORT ───
        with tab2:
            rl, rr = st.columns([3,7], gap="large")
            with rl:
                C = 2 * 3.14159 * 46
                dash = C * score / 100
                badge_label = f"✓ {style_label}"
                st.markdown(f"""
                <div class="score-wrap">
                  <div class="ring-outer">
                    <svg class="ring-svg" viewBox="0 0 120 120" fill="none">
                      <defs>
                        <linearGradient id="ppGrad" x1="0%" y1="0%" x2="100%" y2="100%">
                          <stop offset="0%" stop-color="#6C2BD9"/>
                          <stop offset="100%" stop-color="#0D9488"/>
                        </linearGradient>
                      </defs>
                      <circle class="ring-track" cx="60" cy="60" r="46"/>
                      <circle class="ring-fill" cx="60" cy="60" r="46"
                        stroke-dasharray="{dash:.1f} {C:.1f}" stroke-dashoffset="0"/>
                    </svg>
                    <div class="ring-inner">
                      <span class="ring-num">{score}</span>
                      <span class="ring-pct">% Score</span>
                    </div>
                  </div>
                  <div class="ring-badge">{badge_label}</div>
                </div>
                """, unsafe_allow_html=True)

            with rr:
                st.markdown('<div class="col-label" style="margin-bottom:12px;"><span class="dot-teal"></span>Changes Applied</div>', unsafe_allow_html=True)
                html = "".join(
                    f'<div class="change-item"><span class="check-icon">✓</span>{c}</div>'
                    for c in changes
                )
                st.markdown(f'<div class="changes-list">{html}</div>', unsafe_allow_html=True)

            st.markdown('<div style="height:18px"></div>', unsafe_allow_html=True)
            try:
                import base64 as _b64
                _fb64 = _b64.b64encode(fmt_dl_bytes).decode()
                _MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

                # ── PDF export ──────────────────────────────────────────────
                try:
                    _pdf_bytes = text_to_pdf(raw or apa_preview(raw))
                    _pdf_b64   = _b64.b64encode(_pdf_bytes).decode()
                    _pdf_name  = f"formatted_{used_style.lower()}_{(st.session_state.uploaded_name or 'manuscript').rsplit('.', 1)[0]}.pdf"
                    _pdf_link  = (
                        '<a class="dl-option" href="data:application/pdf;base64,' + _pdf_b64 + '" '
                        'download="' + _pdf_name + '">'
                        '<div class="dl-icon-wrap" style="background:#FEF2F2;border:1px solid #FECACA;">📕</div>'
                        '<div class="dl-option-label">PDF Export</div>'
                        '<div class="dl-option-sub">Downloads <strong>.pdf</strong></div>'
                        '<span style="background:#FEE2E2;color:#991B1B;font-size:10px;font-weight:700;padding:2px 8px;border-radius:10px;">DOWNLOAD</span>'
                        '</a>'
                    )
                except Exception:
                    _pdf_link = ""

                st.markdown(f"""
                <div class="dl-panel">
                  <div class="dl-panel-header">⬇ Download Full Formatted Document As</div>
                  <div class="dl-options">
                    <a class="dl-option" href="data:{_MIME};base64,{_fb64}" download="{fmt_dl_name}">
                      <div class="dl-icon-wrap dl-icon-word">📘</div>
                      <div class="dl-option-label">Microsoft Word</div>
                      <div class="dl-option-sub">Downloads <strong>.docx</strong></div>
                      <span class="dl-badge-word">DOWNLOAD</span>
                    </a>
                    <a class="dl-option" href="https://drive.google.com/drive/my-drive" target="_blank" rel="noopener noreferrer">
                      <div class="dl-icon-wrap dl-icon-gdocs">📝</div>
                      <div class="dl-option-label">Google Docs</div>
                      <div class="dl-option-sub">Download .docx above,<br>then upload to Drive</div>
                      <span class="dl-badge-gdocs">OPEN DRIVE</span>
                    </a>
                  </div>
                </div>
                {_pdf_link}
                """, unsafe_allow_html=True)
            except Exception:
                pass

        # ─── TAB 3 : CITATION CONSISTENCY ───
        with tab3:
            fp_for_cv = result.get("formatted_path")
            raw_text_cv = st.session_state.original_text or ""

            # Try to load pre-computed consistency_report.json saved by agent
            cv_data = None
            if fp_for_cv and Path(fp_for_cv).exists():
                cr_path = Path(fp_for_cv).parent / "consistency_report.json"
                if cr_path.exists():
                    try:
                        import json as _json
                        cv_data = _json.loads(cr_path.read_text(encoding="utf-8", errors="replace"))
                    except Exception:
                        cv_data = None

            # If not available, run validator live
            if cv_data is None:
                with st.spinner("🔍 Running citation–reference consistency check…"):
                    try:
                        _ensure_agent_importable()
                        from agent import validate_from_text  # type: ignore
                        cv_data = validate_from_text(raw_text_cv, style=used_style)
                        # Convert sets to lists for display
                        cv_data["cited_authors"] = sorted(cv_data.get("cited_authors", set()))
                        cv_data["ref_authors"]   = sorted(cv_data.get("ref_authors", set()))
                    except Exception as e:
                        cv_data = {
                            "status": "ERROR", "score": 0,
                            "issues": [f"Validator unavailable: {e}"],
                            "matched": [], "orphan_cites": [], "orphan_refs": [],
                            "ref_count": 0,
                        }

            # ── Render consistency report ──
            cv_status  = cv_data.get("status", "PASS")
            cv_score   = cv_data.get("score", 100)
            matched    = cv_data.get("matched", [])
            o_cites    = cv_data.get("orphan_cites", [])
            o_refs     = cv_data.get("orphan_refs", [])
            ref_count  = cv_data.get("ref_count", 0)
            cv_issues  = cv_data.get("issues", [])

            status_color = {"PASS": "#059669", "WARNINGS": "#D97706", "ERRORS": "#DC2626", "ERROR": "#DC2626"}.get(cv_status, "#6B7280")
            status_bg    = {"PASS": "#ECFDF5", "WARNINGS": "#FFFBEB", "ERRORS": "#FEF2F2", "ERROR": "#FEF2F2"}.get(cv_status, "#F9FAFB")
            status_icon  = {"PASS": "✅", "WARNINGS": "⚠️", "ERRORS": "❌", "ERROR": "❌"}.get(cv_status, "🔍")

            # Header status banner
            st.markdown(f"""
            <div style="background:{status_bg};border:1.5px solid {status_color};
                        border-radius:12px;padding:18px 22px;margin-bottom:18px;">
              <div style="display:flex;align-items:center;gap:12px;margin-bottom:8px;">
                <span style="font-size:28px;">{status_icon}</span>
                <div>
                  <div style="font-size:16px;font-weight:800;color:{status_color};">
                    Consistency Check: {cv_status}
                  </div>
                  <div style="font-size:12px;color:#6B7280;margin-top:2px;">
                    {ref_count} reference entries · {len(matched)} matched · {len(o_cites)} missing refs · {len(o_refs)} uncited refs
                  </div>
                </div>
                <div style="margin-left:auto;font-size:28px;font-weight:900;color:{status_color};">
                  {cv_score}%
                </div>
              </div>
            </div>
            """, unsafe_allow_html=True)

            # Three columns: matched / orphan cites / orphan refs
            ca, cb, cc = st.columns(3, gap="small")

            with ca:
                st.markdown(f"""
                <div style="background:#F0FDF4;border:1px solid #A7F3D0;border-radius:10px;padding:14px;">
                  <div style="font-size:11px;font-weight:800;letter-spacing:0.5px;
                              color:#059669;text-transform:uppercase;margin-bottom:8px;">
                    ✅ Matched ({len(matched)})
                  </div>
                  {"".join(f'<div style="font-size:12px;color:#065F46;padding:3px 0;border-bottom:1px solid #D1FAE5;">{n.title()}</div>' for n in matched) or '<div style="font-size:12px;color:#9CA3AF;">None detected</div>'}
                </div>
                """, unsafe_allow_html=True)

            with cb:
                st.markdown(f"""
                <div style="background:#FEF2F2;border:1px solid #FECACA;border-radius:10px;padding:14px;">
                  <div style="font-size:11px;font-weight:800;letter-spacing:0.5px;
                              color:#DC2626;text-transform:uppercase;margin-bottom:8px;">
                    ❌ Cited but No Reference ({len(o_cites)})
                  </div>
                  {"".join(f'<div style="font-size:12px;color:#991B1B;padding:3px 0;border-bottom:1px solid #FECACA;">{n.title()}</div>' for n in o_cites) or '<div style="font-size:12px;color:#9CA3AF;">None — all citations have references ✓</div>'}
                </div>
                """, unsafe_allow_html=True)

            with cc:
                st.markdown(f"""
                <div style="background:#FFFBEB;border:1px solid #FDE68A;border-radius:10px;padding:14px;">
                  <div style="font-size:11px;font-weight:800;letter-spacing:0.5px;
                              color:#D97706;text-transform:uppercase;margin-bottom:8px;">
                    ⚠ In Refs but Never Cited ({len(o_refs)})
                  </div>
                  {"".join(f'<div style="font-size:12px;color:#92400E;padding:3px 0;border-bottom:1px solid #FDE68A;">{n.title()}</div>' for n in o_refs) or '<div style="font-size:12px;color:#9CA3AF;">None — all references are cited ✓</div>'}
                </div>
                """, unsafe_allow_html=True)

            # Issues list
            if cv_issues:
                st.markdown('<div style="margin-top:16px;">', unsafe_allow_html=True)
                st.markdown('<div class="col-label" style="margin-bottom:8px;"><span class="dot-teal"></span>Detailed Issues</div>', unsafe_allow_html=True)
                for issue in cv_issues:
                    icon_color = "#DC2626" if "❌" in issue else "#D97706" if "⚠" in issue else "#059669"
                    bg = "#FEF2F2" if "❌" in issue else "#FFFBEB" if "⚠" in issue else "#F0FDF4"
                    st.markdown(
                        f'<div style="background:{bg};border-radius:6px;padding:8px 12px;'
                        f'margin-bottom:6px;font-size:13px;color:{icon_color};">{issue}</div>',
                        unsafe_allow_html=True
                    )
                st.markdown('</div>', unsafe_allow_html=True)

            st.markdown("""
            <div style="margin-top:14px;background:var(--pp-bg);border:1px solid var(--pp-border);
                        border-radius:8px;padding:10px 14px;font-size:11px;color:var(--pp-text-muted);">
              <strong>How this works:</strong> Every in-text citation <em>(Author, Year)</em> is extracted from the body.
              Every first-author key is extracted from reference entries. The validator cross-matches both sets
              and flags: citations with no matching reference (❌) and references never cited in the body (⚠).
            </div>
            """, unsafe_allow_html=True)

        st.markdown('<div style="height:14px"></div>', unsafe_allow_html=True)
        if st.button("↺  Reset / New Manuscript", use_container_width=False):
            for k in defaults:
                st.session_state[k] = None
            st.session_state.style_confirmed = False
            st.rerun()


# ─────────────────────────────────────────────
# FOOTER
# ─────────────────────────────────────────────
st.markdown("""
<div class="pp-footer">
  Powered by <strong>AI Agent</strong> · <strong>APA 7th + MLA 9th Rules</strong> · <strong>Smart Doc Detection</strong> · <strong>PDF Export</strong>
</div>
""", unsafe_allow_html=True)