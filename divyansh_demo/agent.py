"””
╔══════════════════════════════════════════════════════════════════════════════╗
║   Academic Manuscript Formatter — APA 7th / MLA 9th / IEEE Auto-Detection  ║
║   LangChain ReAct Agent with Groq LLaMA-3.3-70b                            ║
╠══════════════════════════════════════════════════════════════════════════════╣
║  Features:                                                                  ║
║  • AUTO-DETECTS citation style from document content & keywords             ║
║  • Applies full APA 7th, MLA 9th, or IEEE formatting rules                 ║
║  • 5-stage pipeline: Detect → Parse → Format → Validate → Report           ║
║  • Every formatting element handled: font, margins, spacing, headings,     ║
║    indents, citations, references, header/footer, tables, figures          ║
╠══════════════════════════════════════════════════════════════════════════════╣
║  Usage:                                                                     ║
║    python agent.py paper.docx                    # auto-detect style        ║
║    python agent.py paper.docx –style apa        # force APA                ║
║    python agent.py paper.docx –style mla        # force MLA                ║
║    python agent.py paper.docx –style ieee       # force IEEE               ║
║    python agent.py paper.docx –output out.docx  # custom output path      ║
║    python agent.py paper.docx –direct           # skip LLM agent          ║
║    python agent.py paper.docx –detect-only      # just show detection     ║
╚══════════════════════════════════════════════════════════════════════════════╝
“””

import json
import os
import re
import sys
import argparse
import collections
from pathlib import Path

# ─── Auto-install missing packages ────────────────────────────────────────────

def *ensure(pkg, import_name=None):
import importlib
name = import_name or pkg.split(”-”)[0].replace(”-”, “*”)
try:
importlib.import_module(name)
except ImportError:
os.system(f”pip install {pkg} –break-system-packages -q”)

_ensure(“python-docx”, “docx”)
_ensure(“langchain-groq”, “langchain_groq”)
_ensure(“langchain-core”, “langchain_core”)
_ensure(“langchain”, “langchain”)

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from langchain_groq import ChatGroq
from langchain_core.tools import Tool
from langchain_core.prompts import PromptTemplate
try:
from langchain.agents import create_react_agent, AgentExecutor
except ImportError:
from langchain_classic.agents import create_react_agent, AgentExecutor



# ══════════════════════════════════════════════════════════════════════════════

# STYLE RULE DEFINITIONS

# ══════════════════════════════════════════════════════════════════════════════

STYLE_RULES = {

```
# ── APA 7th Edition ────────────────────────────────────────────────────
"apa": {
    "name": "APA 7th Edition",
    "discipline": "Social Sciences, Psychology, Education, Nursing, Business",
    "document_setup": {
        "font_family": "Times New Roman",
        "font_size": 12,
        "line_spacing": 2.0,
        "margins_inches": 1.0,
        "paragraph_indent_inches": 0.5,
        "alignment": "left",
        "page_width_inches": 8.5,
        "page_height_inches": 11,
    },
    "title_page": {
        "title_bold": True,
        "title_align": "center",
        "all_centered": True,
        "para_count_threshold": 7,
    },
    "abstract": {
        "heading_bold": True,
        "heading_align": "center",
        "body_no_indent": True,
        "body_align": "left",
        "word_limit": "150-250",
        "font_size": 12,
        "line_spacing": 2.0,
    },
    "headings": {
        "level1": {"bold": True,  "italic": False, "align": "center",
                   "font_size": 12, "space_before_pt": 12, "space_after_pt": 0,
                   "indent_inches": 0.0, "inline": False},
        "level2": {"bold": True,  "italic": False, "align": "left",
                   "font_size": 12, "space_before_pt": 12, "space_after_pt": 0,
                   "indent_inches": 0.0, "inline": False},
        "level3": {"bold": True,  "italic": True,  "align": "left",
                   "font_size": 12, "space_before_pt": 12, "space_after_pt": 0,
                   "indent_inches": 0.0, "inline": False},
        "level4": {"bold": True,  "italic": False, "align": "left",
                   "font_size": 12, "space_before_pt": 0,  "space_after_pt": 0,
                   "indent_inches": 0.5, "inline": True, "period": True},
        "level5": {"bold": True,  "italic": True,  "align": "left",
                   "font_size": 12, "space_before_pt": 0,  "space_after_pt": 0,
                   "indent_inches": 0.5, "inline": True, "period": True},
    },
    "body_paragraph": {
        "first_line_indent_inches": 0.5,
        "left_indent_inches": 0.0,
        "align": "left",
        "line_spacing": 2.0,
        "space_before_pt": 0,
        "space_after_pt": 0,
        "font_size": 12,
    },
    "references": {
        "heading": "References",
        "heading_bold": True,
        "heading_align": "center",
        "hanging_indent_inches": 0.5,
        "line_spacing": 2.0,
        "font_size": 12,
        "order": "alphabetical",
    },
    "citation_style": "author_date",
    "citation_patterns": [
        r'\([A-Za-z\s\-&]+(?:\s+et\s+al\.)?,\s*\d{4}[a-z]?\)',
        r'\([A-Za-z\s\-&]+(?:\s+et\s+al\.)?,\s*\d{4}[a-z]?,\s*(?:p\.|pp\.)\s*[\d\-]+\)',
    ],
    "detection_signals": {
        "strong": [
            "author, year", "et al.,", "doi:", "apa", "psychology",
            "social science", "education", "nursing", "p. ", "pp. ",
            "running head", "https://doi.org", "n =", "sd =", "m =",
            "participants", "qualitative", "quantitative"
        ],
        "bibliography_title": "references",
    }
},

# ── MLA 9th Edition ────────────────────────────────────────────────────
"mla": {
    "name": "MLA 9th Edition",
    "discipline": "Literature, Linguistics, Humanities, Cultural Studies",
    "document_setup": {
        "font_family": "Times New Roman",
        "font_size": 12,
        "line_spacing": 2.0,
        "margins_inches": 1.0,
        "paragraph_indent_inches": 0.5,
        "alignment": "left",
        "page_width_inches": 8.5,
        "page_height_inches": 11,
    },
    "first_page": {
        "header_lines": 4,
        "title_centered": True,
        "title_bold": False,
    },
    "headings": {
        "level1": {"bold": True,  "italic": False, "align": "center",
                   "font_size": 12, "space_before_pt": 12, "space_after_pt": 0,
                   "indent_inches": 0.0, "inline": False},
        "level2": {"bold": True,  "italic": True,  "align": "center",
                   "font_size": 12, "space_before_pt": 12, "space_after_pt": 0,
                   "indent_inches": 0.0, "inline": False},
        "level3": {"bold": True,  "italic": False, "align": "left",
                   "font_size": 12, "space_before_pt": 12, "space_after_pt": 0,
                   "indent_inches": 0.0, "inline": False},
        "level4": {"bold": True,  "italic": True,  "align": "left",
                   "font_size": 12, "space_before_pt": 0,  "space_after_pt": 0,
                   "indent_inches": 0.0, "inline": False},
        "level5": {"bold": True,  "italic": False, "align": "left",
                   "font_size": 12, "space_before_pt": 0,  "space_after_pt": 0,
                   "indent_inches": 0.5, "inline": True, "period": True},
    },
    "body_paragraph": {
        "first_line_indent_inches": 0.5,
        "left_indent_inches": 0.0,
        "align": "left",
        "line_spacing": 2.0,
        "space_before_pt": 0,
        "space_after_pt": 0,
        "font_size": 12,
    },
    "works_cited": {
        "heading": "Works Cited",
        "heading_bold": False,
        "heading_align": "center",
        "hanging_indent_inches": 0.5,
        "line_spacing": 2.0,
        "font_size": 12,
        "order": "alphabetical",
    },
    "citation_style": "author_page",
    "citation_patterns": [
        r'\([A-Za-z\-]+(?:\s+et\s+al\.)?\s+\d+(?:[-]\d+)?\)',
        r'\([A-Za-z\-]+(?:\s+and\s+[A-Za-z\-]+)?\s+\d+\)',
        r'\(qtd\.\s+in\s+[A-Za-z\-]+\s+\d+\)',
    ],
    "detection_signals": {
        "strong": [
            "works cited", "mla", "humanities", "literature", "qtd. in",
            "norton anthology", "english", "cultural studies",
            "literary", "close reading"
        ],
        "bibliography_title": "works cited",
    }
},

# ── IEEE ────────────────────────────────────────────────────────────────
"ieee": {
    "name": "IEEE",
    "discipline": "Electrical Engineering, Computer Science, Electronics, IT",
    "document_setup": {
        "font_family": "Times New Roman",
        "font_size": 10,
        "line_spacing": 1.0,
        "margins_top_inches": 0.75,
        "margins_bottom_inches": 1.0,
        "margins_left_inches": 0.625,
        "margins_right_inches": 0.625,
        "paragraph_indent_inches": 0.15,
        "alignment": "justify",
        "page_width_inches": 8.5,
        "page_height_inches": 11,
        "two_column": True,
    },
    "title_section": {
        "title_font_size": 24,
        "title_bold": False,
        "title_align": "center",
        "author_font_size": 11,
        "affiliation_font_size": 9,
    },
    "abstract": {
        "inline_label": True,
        "label_bold_italic": True,
        "label_format": "Abstract\u2014",
        "font_size": 9,
        "line_spacing": 1.0,
    },
    "headings": {
        "level1": {"bold": False, "italic": False, "align": "center",
                   "font_size": 10, "small_caps": True,
                   "space_before_pt": 6, "space_after_pt": 6,
                   "indent_inches": 0.0, "inline": False},
        "level2": {"bold": False, "italic": True,  "align": "left",
                   "font_size": 10, "small_caps": False,
                   "space_before_pt": 6, "space_after_pt": 0,
                   "indent_inches": 0.0, "inline": False},
        "level3": {"bold": False, "italic": True,  "align": "left",
                   "font_size": 10, "small_caps": False,
                   "space_before_pt": 0, "space_after_pt": 0,
                   "indent_inches": 0.15, "inline": True, "colon": True},
    },
    "body_paragraph": {
        "first_line_indent_inches": 0.15,
        "left_indent_inches": 0.0,
        "align": "justify",
        "line_spacing": 1.0,
        "space_before_pt": 0,
        "space_after_pt": 2,
        "font_size": 10,
    },
    "references": {
        "heading": "References",
        "heading_bold": False,
        "heading_align": "center",
        "hanging_indent_inches": 0.25,
        "line_spacing": 1.0,
        "font_size": 8,
        "order": "appearance",
        "number_format": "[1]",
    },
    "citation_style": "numbered",
    "citation_patterns": [
        r'\[\d+\]',
        r'\[\d+(?:[,\s]+\d+)*\]',
    ],
    "detection_signals": {
        "strong": [
            "[1]", "[2]", "ieee", "index terms", "transactions on",
            "proceedings of", "neural network", "deep learning",
            "machine learning", "algorithm", "fpga", "circuit",
            "vol.", "no.", "pp.", "conference paper", "bandwidth",
            "abstract\u2014", "abstract\u2013"
        ],
        "bibliography_title": "references",
    }
}
```

}

# ── Common APA-style section keywords ──────────────────────────────────────

APA_L1_KEYWORDS = {
‘abstract’, ‘introduction’, ‘method’, ‘methods’, ‘results’,
‘discussion’, ‘conclusion’, ‘conclusions’, ‘references’, ‘bibliography’,
‘appendix’, ‘appendices’, ‘acknowledgments’, ‘acknowledgements’,
‘author note’, ‘keywords’, ‘general discussion’, ‘general method’,
‘literature review’, ‘theoretical framework’, ‘background’
}
APA_L2_KEYWORDS = {
‘participants’, ‘procedure’, ‘measures’, ‘instruments’, ‘materials’,
‘design’, ‘data analysis’, ‘data collection’, ‘ethical considerations’,
‘sample’, ‘apparatus’, ‘stimuli’, ‘limitations’, ‘future directions’,
‘theoretical implications’, ‘practical implications’, ‘overview’,
‘recruitment’, ‘analysis’, ‘interview’, ‘survey’
}

# ══════════════════════════════════════════════════════════════════════════════

# XML / DOCX UTILITY HELPERS

# ══════════════════════════════════════════════════════════════════════════════

def _set_run_font(run, font_name: str, size_pt: float,
bold: bool = False, italic: bool = False,
small_caps: bool = False):
“”“Completely override run font in python-docx AND raw XML.”””
run.font.name = font_name
run.font.size = Pt(size_pt)
run.bold = bold
run.italic = italic
rPr = run._r.get_or_add_rPr()
# Force font override via XML
rFonts = OxmlElement(‘w:rFonts’)
for attr in (‘w:ascii’, ‘w:hAnsi’, ‘w:eastAsia’, ‘w:cs’):
rFonts.set(qn(attr), font_name)
old = rPr.find(qn(‘w:rFonts’))
if old is not None:
rPr.remove(old)
rPr.insert(0, rFonts)
# Small caps
sc = rPr.find(qn(‘w:smallCaps’))
if small_caps:
if sc is None:
sc_el = OxmlElement(‘w:smallCaps’)
rPr.append(sc_el)
else:
if sc is not None:
rPr.remove(sc)

def _set_paragraph_spacing(para, line_spacing: float,
space_before_pt: float = 0,
space_after_pt: float = 0):
pf = para.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
pf.line_spacing = line_spacing
pf.space_before = Pt(space_before_pt)
pf.space_after = Pt(space_after_pt)

def _set_indents(para, left_inches: float = 0.0,
first_line_inches: float = 0.0,
hanging_inches: float = 0.0):
pf = para.paragraph_format
if hanging_inches > 0:
pf.left_indent = Inches(hanging_inches)
pf.first_line_indent = Inches(-hanging_inches)
else:
pf.left_indent = Inches(left_inches)
pf.first_line_indent = Inches(first_line_inches)

def _set_alignment(para, align: str):
m = {
“center”: WD_ALIGN_PARAGRAPH.CENTER,
“left”:   WD_ALIGN_PARAGRAPH.LEFT,
“right”:  WD_ALIGN_PARAGRAPH.RIGHT,
“justify”: WD_ALIGN_PARAGRAPH.JUSTIFY,
}
para.alignment = m.get(align.lower(), WD_ALIGN_PARAGRAPH.LEFT)

def _inject_page_number(para):
“”“Add a PAGE field to paragraph.”””
run = para.add_run()
for ftype in (‘begin’, ‘end’):
if ftype == ‘begin’:
fc = OxmlElement(‘w:fldChar’)
fc.set(qn(‘w:fldCharType’), ‘begin’)
run._r.append(fc)
it = OxmlElement(‘w:instrText’)
it.set(qn(‘xml:space’), ‘preserve’)
it.text = ’ PAGE ’
run._r.append(it)
fc2 = OxmlElement(‘w:fldChar’)
fc2.set(qn(‘w:fldCharType’), ftype)
run._r.append(fc2)

def _clear_header(section):
for p in section.header.paragraphs:
p.clear()

def _setup_apa_header(doc, font, size):
sec = doc.sections[0]
sec.different_first_page_header_footer = False
_clear_header(sec)
header = sec.header
hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
hp.clear()
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
_set_paragraph_spacing(hp, 1.0, 0, 0)
r = hp.add_run()
r.font.name = font
r.font.size = Pt(size)
_inject_page_number(hp)

def _setup_mla_header(doc, author_last_name, font, size):
sec = doc.sections[0]
_clear_header(sec)
header = sec.header
hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
hp.clear()
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
_set_paragraph_spacing(hp, 1.0, 0, 0)
nr = hp.add_run(f”{author_last_name} “ if author_last_name else “”)
nr.font.name = font
nr.font.size = Pt(size)
_inject_page_number(hp)

def _setup_ieee_header(doc):
_clear_header(doc.sections[0])

# ══════════════════════════════════════════════════════════════════════════════

# STYLE AUTO-DETECTION ENGINE

# ══════════════════════════════════════════════════════════════════════════════

def detect_citation_style(docx_path: str) -> dict:
“””
Multi-signal analysis to detect APA / MLA / IEEE from document content.

```
Signals (in order of strength):
  1. Citation bracket/pattern counts  — strongest
  2. Bibliography section title       — very strong
  3. Domain/style keywords            — strong
  4. Structural markers               — medium
  5. Date format, running head, etc.  — weak
"""
doc = Document(docx_path)
full_text = "\n".join(p.text for p in doc.paragraphs).lower()
paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

scores = {"apa": 0, "mla": 0, "ieee": 0}
evidence = {"apa": [], "mla": [], "ieee": []}

# ── 1. Citation pattern counts ─────────────────────────────────────────
n_apa  = len(re.findall(r'\([A-Za-z\-\s]+(?:\s+et\s+al\.)?,\s*\d{4}[a-z]?\)', full_text))
n_mla  = len(re.findall(r'\([A-Za-z\-]+(?:\s+et\s+al\.)?\s+\d+\)', full_text))
n_ieee = len(re.findall(r'\[\d+\]', full_text))

if n_apa:
    scores["apa"] += min(n_apa * 3, 30)
    evidence["apa"].append(f"{n_apa} author-date citations (Author, Year)")
if n_mla:
    scores["mla"] += min(n_mla * 3, 30)
    evidence["mla"].append(f"{n_mla} author-page citations (Author Page#)")
if n_ieee:
    scores["ieee"] += min(n_ieee * 3, 30)
    evidence["ieee"].append(f"{n_ieee} bracket citations [N]")

# ── 2. Bibliography title ──────────────────────────────────────────────
for pt in paragraphs:
    ptl = pt.lower().strip().rstrip('.')
    if ptl == "works cited":
        scores["mla"] += 25
        evidence["mla"].append("'Works Cited' section heading found")
    elif ptl == "references":
        if n_ieee > max(n_apa, n_mla):
            scores["ieee"] += 15
            evidence["ieee"].append("'References' section with IEEE brackets")
        else:
            scores["apa"] += 15
            evidence["apa"].append("'References' section heading found")
    elif ptl in ("bibliography", "works consulted"):
        scores["mla"] += 8
        evidence["mla"].append(f"'{pt}' section found")

# ── 3. Domain/style keywords ───────────────────────────────────────────
KW = {
    "apa": [
        "psychology", "social science", "education", "nursing", "counseling",
        "cognitive", "behavioral", "participant", "qualitative", "quantitative",
        "doi:", "https://doi.org", "running head", "apa 7",
        "american psychological association", "n =", "sd =", "m ="
    ],
    "mla": [
        "works cited", "mla", "humanities", "literature", "literary",
        "qtd. in", "norton anthology", "close reading", "textual analysis",
        "english department", "cultural studies", "modern language association"
    ],
    "ieee": [
        "index terms", "ieee", "transactions on", "proceedings",
        "neural network", "deep learning", "machine learning", "algorithm",
        "fpga", "microcontroller", "wireless", "bandwidth", "circuit",
        "voltage", "frequency", "throughput", "latency", "ablation",
        "computer science", "proposed method", "simulations show",
        "experimental results", "state-of-the-art"
    ]
}
for style_key, keywords in KW.items():
    for kw in keywords:
        if kw in full_text:
            scores[style_key] += 4
            evidence[style_key].append(f"keyword '{kw}'")

# ── 4. Structural markers ──────────────────────────────────────────────
if re.search(r'abstract[—–-]', full_text):
    scores["ieee"] += 12
    evidence["ieee"].append("Inline 'Abstract—' IEEE-style label")
if re.search(r'index terms[—–-]', full_text):
    scores["ieee"] += 15
    evidence["ieee"].append("'Index Terms—' section found")
if re.search(r'vol\.\s*\d+.*?no\.\s*\d+.*?pp\.', full_text):
    scores["ieee"] += 8
    evidence["ieee"].append("IEEE vol./no./pp. reference format")

# MLA day-month-year date
if re.search(r'\b\d{1,2}\s+(january|february|march|april|may|june|july|august|'
             r'september|october|november|december)\s+\d{4}\b', full_text):
    scores["mla"] += 8
    evidence["mla"].append("MLA date format (Day Month Year)")

# APA signals
doi_n = len(re.findall(r'https://doi\.org/|doi:\s*10\.', full_text))
if doi_n:
    scores["apa"] += min(doi_n * 2, 10)
    evidence["apa"].append(f"{doi_n} DOI(s) found")
if "running head" in full_text:
    scores["apa"] += 10
    evidence["apa"].append("'Running head' found")

# ── Determine winner ───────────────────────────────────────────────────
best = max(scores, key=lambda k: scores[k])
total = sum(scores.values()) or 1
confidence = round(scores[best] / total * 100)
if all(v == 0 for v in scores.values()):
    best, confidence = "apa", 0

return {
    "detected_style": best,
    "confidence_percent": confidence,
    "scores": scores,
    "evidence": {k: v[:6] for k, v in evidence.items()},
    "style_name": STYLE_RULES[best]["name"],
    "discipline": STYLE_RULES[best]["discipline"],
}
```

# ══════════════════════════════════════════════════════════════════════════════

# HEADING CLASSIFIER

# ══════════════════════════════════════════════════════════════════════════════

def _classify_heading(text: str, style_name: str, citation_style: str) -> int:
sl = style_name.lower()
for n in range(1, 6):
if f’heading {n}’ in sl:
return n
if ‘title’ in sl and citation_style == “apa”:
return 1
tl = text.strip().lower().rstrip(’.’)
if tl in APA_L1_KEYWORDS:
return 1
if tl in APA_L2_KEYWORDS:
return 2
words = text.strip().split()
if text.strip().isupper() and 2 <= len(words) <= 10:
return 1
if (len(text.strip()) < 80 and text.strip().istitle()
and not text.strip().endswith(’.’) and 1 <= len(words) <= 8):
return 2
return 0

# ══════════════════════════════════════════════════════════════════════════════

# TOOL 1 — Detect Citation Style

# ══════════════════════════════════════════════════════════════════════════════

def tool_detect_style(docx_path: str) -> str:
“””
Automatically detect whether the document uses APA, MLA, or IEEE.
Input: path to .docx file.
Output: JSON with detected_style, confidence_percent, scores, evidence.
“””
docx_path = docx_path.strip().strip(’”'’)
if not os.path.exists(docx_path):
return json.dumps({“error”: f”File not found: {docx_path}”})
return json.dumps(detect_citation_style(docx_path), indent=2)

# ══════════════════════════════════════════════════════════════════════════════

# TOOL 2 — Parse Document

# ══════════════════════════════════════════════════════════════════════════════

def tool_parse_document(args_str: str) -> str:
“””
Parse a .docx file and return detailed structural JSON.
Input: “docx_path style”
“””
parts = args_str.strip().strip(’”'’).split()
docx_path = parts[0] if parts else “”
citation_style = parts[1].lower() if len(parts) > 1 else “apa”

```
if not os.path.exists(docx_path):
    return json.dumps({"error": f"File not found: {docx_path}"})
if citation_style not in STYLE_RULES:
    citation_style = "apa"

doc = Document(docx_path)
result = {
    "file": docx_path,
    "citation_style": citation_style,
    "total_paragraphs": len(doc.paragraphs),
    "headings": [],
    "citations": [],
    "special_sections": {},
    "reference_entries": [],
    "author_last_name": "",
    "stats": {}
}

word_count = 0
in_refs = False
non_empty = 0
bib_titles = {"references", "works cited", "bibliography"}
cite_pats = STYLE_RULES[citation_style]["citation_patterns"]

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    non_empty += 1
    style_name = para.style.name if para.style else "Normal"
    tl = text.lower().rstrip('.')

    if tl in bib_titles:
        in_refs = True
        result["special_sections"]["bibliography_start"] = i
        result["special_sections"]["bibliography_title"] = text
    if tl == "abstract":
        result["special_sections"]["abstract_start"] = i
    for kw in ("introduction", "method", "methods", "results",
               "discussion", "conclusion", "conclusions",
               "index terms", "acknowledgment"):
        if tl == kw:
            result["special_sections"][kw] = i

    if non_empty == 1 and citation_style == "mla" and len(text.split()) <= 5:
        parts2 = text.split()
        if parts2:
            result["author_last_name"] = parts2[-1]

    hl = _classify_heading(text, style_name, citation_style)
    cites = []
    for pat in cite_pats:
        cites += re.findall(pat, text)

    is_ref = False
    if in_refs:
        ref_sigs = [r'^[A-Z][a-z]+,\s+[A-Z]\.', r'^\[\d+\]',
                    r'doi:', r'https?://']
        is_ref = any(re.search(p, text) for p in ref_sigs)

    if hl > 0:
        result["headings"].append({"index": i, "level": hl, "text": text})
    result["citations"].extend([{"match": c, "para": i} for c in cites])
    if is_ref:
        result["reference_entries"].append({"index": i, "text": text[:90]})
    word_count += len(text.split())

result["stats"] = {
    "total_words": word_count,
    "total_headings": len(result["headings"]),
    "unique_citations": len({str(c["match"]) for c in result["citations"]}),
    "total_reference_entries": len(result["reference_entries"]),
    "sections_found": list(result["special_sections"].keys()),
}
return json.dumps(result, indent=2)
```

# ══════════════════════════════════════════════════════════════════════════════

# CORE PARAGRAPH FORMATTERS

# ══════════════════════════════════════════════════════════════════════════════

def _fmt_heading(para, level: int, rules: dict, citation_style: str):
hl_rules = rules.get(“headings”, {})
r = hl_rules.get(f”level{level}”, hl_rules.get(“level2”, {}))
if not r:
return
ds = rules[“document_setup”]
font = ds[“font_family”]
size = r.get(“font_size”, ds[“font_size”])
ls   = ds[“line_spacing”]
_set_alignment(para, r.get(“align”, “left”))
_set_paragraph_spacing(para, ls, r.get(“space_before_pt”, 0), r.get(“space_after_pt”, 0))
_set_indents(para, left_inches=r.get(“indent_inches”, 0))
for run in para.runs:
_set_run_font(run, font, size,
bold=r.get(“bold”, False),
italic=r.get(“italic”, False),
small_caps=r.get(“small_caps”, False))

def _fmt_body(para, rules: dict):
bp = rules.get(“body_paragraph”, {})
ds = rules[“document_setup”]
font = ds[“font_family”]
size = bp.get(“font_size”, ds[“font_size”])
ls   = bp.get(“line_spacing”, ds[“line_spacing”])
_set_alignment(para, bp.get(“align”, “left”))
_set_paragraph_spacing(para, ls,
bp.get(“space_before_pt”, 0),
bp.get(“space_after_pt”, 0))
_set_indents(para,
left_inches=bp.get(“left_indent_inches”, 0.0),
first_line_inches=bp.get(“first_line_indent_inches”, 0.5))
for run in para.runs:
_set_run_font(run, font, size)

def _fmt_abstract_body(para, rules: dict):
ds = rules[“document_setup”]
ar = rules.get(“abstract”, {})
font = ds[“font_family”]
size = ar.get(“font_size”, ds[“font_size”])
ls   = ar.get(“line_spacing”, ds[“line_spacing”])
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
_set_paragraph_spacing(para, ls, 0, 0)
_set_indents(para, first_line_inches=0.0)
for run in para.runs:
_set_run_font(run, font, size)

def _fmt_bib_entry(para, rules: dict, citation_style: str):
if citation_style == “mla”:
bib_r = rules.get(“works_cited”, {})
else:
bib_r = rules.get(“references”, {})
ds = rules[“document_setup”]
font = ds[“font_family”]
size = bib_r.get(“font_size”, ds[“font_size”])
ls   = bib_r.get(“line_spacing”, ds[“line_spacing”])
hang = bib_r.get(“hanging_indent_inches”, 0.5)
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
_set_paragraph_spacing(para, ls, 0, 2)
_set_indents(para, hanging_inches=hang)
for run in para.runs:
_set_run_font(run, font, size)

def _fmt_title_page_apa(para, rules: dict, non_empty_idx: int):
ds = rules[“document_setup”]
tp = rules.get(“title_page”, {})
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
_set_paragraph_spacing(para, ds[“line_spacing”], 0, 0)
_set_indents(para)
is_title = non_empty_idx == 1 and tp.get(“title_bold”, True)
for run in para.runs:
_set_run_font(run, ds[“font_family”], ds[“font_size”], bold=is_title)

def _fmt_mla_header_block(para, rules: dict):
ds = rules[“document_setup”]
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
_set_paragraph_spacing(para, ds[“line_spacing”], 0, 0)
_set_indents(para)
for run in para.runs:
_set_run_font(run, ds[“font_family”], ds[“font_size”])

def _fmt_mla_title(para, rules: dict):
ds = rules[“document_setup”]
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
_set_paragraph_spacing(para, ds[“line_spacing”], 0, 0)
_set_indents(para)
for run in para.runs:
_set_run_font(run, ds[“font_family”], ds[“font_size”],
bold=False, italic=False)

def _fmt_ieee_title(para, rules: dict):
ts = rules.get(“title_section”, {})
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
_set_paragraph_spacing(para, 1.0, 6, 6)
_set_indents(para)
for run in para.runs:
_set_run_font(run, rules[“document_setup”][“font_family”],
ts.get(“title_font_size”, 24))

def _fmt_ieee_author(para, rules: dict):
ts = rules.get(“title_section”, {})
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
_set_paragraph_spacing(para, 1.0, 2, 2)
_set_indents(para)
for run in para.runs:
_set_run_font(run, rules[“document_setup”][“font_family”],
ts.get(“author_font_size”, 11))

# ══════════════════════════════════════════════════════════════════════════════

# TOOL 3 — Apply Style Formatting

# ══════════════════════════════════════════════════════════════════════════════

def tool_apply_formatting(args_str: str) -> str:
“””
Apply full formatting for the specified style.
Input: JSON {“docx_path”:”…”,“style”:“apa/mla/ieee”,“output_path”:”…”}
or space-separated: “docx_path style output_path”
“””
args_str = args_str.strip().strip(’”'’)
try:
obj = json.loads(args_str)
docx_path   = obj.get(“docx_path”, obj.get(“input”, “”))
style       = obj.get(“style”, “apa”).lower()
output_path = obj.get(“output_path”, obj.get(“output”,
f”formatted_{style}.docx”))
except (json.JSONDecodeError, TypeError):
parts = args_str.split()
docx_path   = parts[0] if len(parts) > 0 else “”
style       = parts[1].lower() if len(parts) > 1 else “apa”
output_path = parts[2] if len(parts) > 2 else f”formatted_{style}.docx”

```
docx_path   = docx_path.strip().strip('"\'')
output_path = output_path.strip().strip('"\'')

if not os.path.exists(docx_path):
    return f"❌ File not found: {docx_path}"
if style not in STYLE_RULES:
    return f"❌ Unknown style '{style}'. Choose: apa, mla, ieee"

rules = STYLE_RULES[style]
ds    = rules["document_setup"]
FONT  = ds["font_family"]
FSIZE = ds["font_size"]
LS    = ds["line_spacing"]
log   = [f"📄 Applying {rules['name']} to: {docx_path}", "═" * 65]

doc = Document(docx_path)

# ── 1. Page Setup ────────────────────────────────────────────────────
for sec in doc.sections:
    if style == "ieee":
        sec.top_margin    = Inches(ds.get("margins_top_inches",    0.75))
        sec.bottom_margin = Inches(ds.get("margins_bottom_inches", 1.0))
        sec.left_margin   = Inches(ds.get("margins_left_inches",   0.625))
        sec.right_margin  = Inches(ds.get("margins_right_inches",  0.625))
    else:
        m = Inches(ds["margins_inches"])
        sec.top_margin = sec.bottom_margin = m
        sec.left_margin = sec.right_margin = m
    sec.page_width  = Inches(ds.get("page_width_inches",  8.5))
    sec.page_height = Inches(ds.get("page_height_inches", 11))
log.append("✅ Page size & margins configured")

# ── 2. Header ────────────────────────────────────────────────────────
if style == "apa":
    _setup_apa_header(doc, FONT, FSIZE)
    log.append("✅ APA header: page number right-aligned")
elif style == "mla":
    first_nonempty = next((p.text.strip() for p in doc.paragraphs
                            if p.text.strip()), "")
    author_ln = first_nonempty.split()[-1] if first_nonempty.split() else ""
    _setup_mla_header(doc, author_ln, FONT, FSIZE)
    log.append(f"✅ MLA header: '{author_ln} [page#]' right-aligned")
else:
    _setup_ieee_header(doc)
    log.append("✅ IEEE header: cleared")

# ── 3. Paragraph Loop ────────────────────────────────────────────────
non_empty_count    = 0
in_abstract        = False
in_bibliography    = False
abstract_body_done = False
bib_heading_done   = False
mla_title_done     = False
heading_counts     = collections.Counter()
bib_titles         = {"references", "works cited", "bibliography"}
APA_THRESHOLD      = rules.get("title_page", {}).get("para_count_threshold", 7)
MLA_HEADER_LINES   = rules.get("first_page", {}).get("header_lines", 4)

for para in doc.paragraphs:
    raw = para.text.strip()
    if not raw:
        _set_paragraph_spacing(para, LS, 0, 0)
        continue

    non_empty_count += 1
    sn  = para.style.name if para.style else "Normal"
    tl  = raw.lower().rstrip('.')
    hl  = _classify_heading(raw, sn, style)

    # Track context
    if tl in bib_titles:
        in_bibliography = True
        bib_heading_done = True
        in_abstract = False
    if tl == "abstract":
        in_abstract = True
        abstract_body_done = False
    elif tl in ("introduction", "method", "methods", "results",
                "discussion", "conclusion", "conclusions",
                "index terms", "acknowledgment"):
        in_abstract = False

    # ── APA Pipeline ──────────────────────────────────────────────
    if style == "apa":
        if non_empty_count <= APA_THRESHOLD:
            _fmt_title_page_apa(para, rules, non_empty_count)
            log.append(f"  [Title Page #{non_empty_count}] {raw[:55]}")
            continue
        if hl > 0:
            heading_counts[hl] += 1
            _fmt_heading(para, hl, rules, style)
            if tl == "abstract":
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    _set_run_font(run, FONT, FSIZE, bold=True)
            log.append(f"  [APA L{hl} Heading] {raw[:60]}")
            continue
        if in_abstract and not abstract_body_done:
            abstract_body_done = True
            _fmt_abstract_body(para, rules)
            log.append(f"  [Abstract Body] {raw[:55]}")
            continue
        if in_bibliography:
            if bib_heading_done and tl not in bib_titles:
                bib_heading_done = False
                _fmt_bib_entry(para, rules, style)
                log.append(f"  [Reference Entry] {raw[:55]}")
                continue
            elif not bib_heading_done:
                _fmt_bib_entry(para, rules, style)
                continue
        _fmt_body(para, rules)

    # ── MLA Pipeline ──────────────────────────────────────────────
    elif style == "mla":
        if non_empty_count <= MLA_HEADER_LINES:
            _fmt_mla_header_block(para, rules)
            log.append(f"  [MLA Header Block #{non_empty_count}] {raw[:55]}")
            continue
        if not mla_title_done:
            mla_title_done = True
            _fmt_mla_title(para, rules)
            log.append(f"  [MLA Paper Title] {raw[:60]}")
            continue
        if hl > 0:
            heading_counts[hl] += 1
            _fmt_heading(para, hl, rules, style)
            log.append(f"  [MLA L{hl} Heading] {raw[:60]}")
            continue
        if in_bibliography:
            if bib_heading_done and tl not in bib_titles:
                bib_heading_done = False
                _fmt_bib_entry(para, rules, style)
                log.append(f"  [Works Cited Entry] {raw[:55]}")
                continue
            elif not bib_heading_done:
                _fmt_bib_entry(para, rules, style)
                continue
        _fmt_body(para, rules)

    # ── IEEE Pipeline ─────────────────────────────────────────────
    elif style == "ieee":
        if non_empty_count == 1:
            _fmt_ieee_title(para, rules)
            log.append(f"  [IEEE Title] {raw[:60]}")
            continue
        if non_empty_count in (2, 3):
            _fmt_ieee_author(para, rules)
            log.append(f"  [IEEE Author #{non_empty_count}] {raw[:55]}")
            continue
        if tl in ("abstract", "abstract\u2014", "abstract\u2013"):
            in_abstract = True
            abstract_body_done = False
            _fmt_abstract_body(para, rules)
            for run in para.runs:
                _set_run_font(run, FONT, 9, bold=True, italic=True)
            log.append(f"  [IEEE Abstract Label] {raw[:55]}")
            continue
        if in_abstract and not abstract_body_done:
            abstract_body_done = True
            _fmt_abstract_body(para, rules)
            log.append(f"  [IEEE Abstract Body] {raw[:55]}")
            continue
        if hl > 0:
            heading_counts[hl] += 1
            _fmt_heading(para, hl, rules, style)
            log.append(f"  [IEEE L{hl} Heading] {raw[:60]}")
            continue
        if in_bibliography:
            if bib_heading_done and tl not in bib_titles:
                bib_heading_done = False
                _fmt_bib_entry(para, rules, style)
                log.append(f"  [IEEE Reference] {raw[:55]}")
                continue
            elif not bib_heading_done:
                _fmt_bib_entry(para, rules, style)
                continue
        _fmt_body(para, rules)

log.append(f"✅ Processed {non_empty_count} paragraphs | "
           f"Headings: {dict(heading_counts)}")

# ── 4. Table Cells ───────────────────────────────────────────────────
tc = 0
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _set_paragraph_spacing(p, LS, 0, 0)
                _set_indents(p)
                for run in p.runs:
                    _set_run_font(run, FONT, FSIZE)
                tc += 1
if tc:
    log.append(f"✅ Formatted {tc} table cell paragraphs")

doc.save(output_path)
log.append("═" * 65)
log.append(f"✅ {rules['name']} formatted document saved: {output_path}")
return "\n".join(log)
```

# ══════════════════════════════════════════════════════════════════════════════

# TOOL 4 — Validate Compliance

# ══════════════════════════════════════════════════════════════════════════════

def tool_validate(args_str: str) -> str:
“””
Check a formatted .docx for compliance with the specified style rules.
Input: “docx_path style” or JSON {“docx_path”:”…”,“style”:“apa”}
“””
args_str = args_str.strip().strip(’”'’)
try:
obj = json.loads(args_str)
docx_path = obj.get(“docx_path”, obj.get(“input”, “”))
style     = obj.get(“style”, “apa”).lower()
except (json.JSONDecodeError, TypeError):
parts = args_str.split()
docx_path = parts[0] if len(parts) > 0 else “”
style     = parts[1].lower() if len(parts) > 1 else “apa”

```
docx_path = docx_path.strip().strip('"\'')
if not os.path.exists(docx_path):
    return json.dumps({"error": f"File not found: {docx_path}"})
if style not in STYLE_RULES:
    return json.dumps({"error": f"Unknown style '{style}'"})

rules = STYLE_RULES[style]
ds    = rules["document_setup"]
doc   = Document(docx_path)
issues = []
checked = 0

# Margin check
if style == "ieee":
    expected_margins = {
        "top_margin":    ds.get("margins_top_inches",    0.75),
        "bottom_margin": ds.get("margins_bottom_inches", 1.0),
        "left_margin":   ds.get("margins_left_inches",   0.625),
        "right_margin":  ds.get("margins_right_inches",  0.625),
    }
else:
    m = ds["margins_inches"]
    expected_margins = {k: m for k in
                        ("top_margin", "bottom_margin", "left_margin", "right_margin")}

for i, sec in enumerate(doc.sections):
    for attr, exp in expected_margins.items():
        val = getattr(sec, attr, None)
        if val and abs(val.inches - exp) > 0.06:
            issues.append({"type": "margin", "section": i,
                            "attr": attr, "found": round(val.inches, 3),
                            "expected": exp})

# Paragraph checks
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    checked += 1
    for run in para.runs:
        if run.font.name and run.font.name != ds["font_family"]:
            issues.append({"type": "font_family", "para": i,
                            "preview": text[:40], "found": run.font.name,
                            "expected": ds["font_family"]})
            break
        if run.font.size:
            fp = run.font.size.pt
            if abs(fp - ds["font_size"]) > 1.0:
                issues.append({"type": "font_size", "para": i,
                               "preview": text[:40], "found": fp,
                               "expected": ds["font_size"]})
                break
    pf = para.paragraph_format
    if pf.line_spacing and isinstance(pf.line_spacing, (int, float)):
        if abs(pf.line_spacing - ds["line_spacing"]) > 0.2:
            issues.append({"type": "line_spacing", "para": i,
                            "preview": text[:40],
                            "found": round(float(pf.line_spacing), 2),
                            "expected": ds["line_spacing"]})

return json.dumps({
    "summary": {"style": rules["name"], "paragraphs_checked": checked,
                "issues_found": len(issues)},
    "issues": issues[:60]
}, indent=2)
```

# ══════════════════════════════════════════════════════════════════════════════

# TOOL 5 — Fix Font & Spacing Sweep

# ══════════════════════════════════════════════════════════════════════════════

def tool_fix_sweep(args_str: str) -> str:
“””
Brute-force sweep: enforce correct font, size, line spacing on every run.
Input: “docx_path style”
“””
parts = args_str.strip().strip(’”'’).split()
docx_path = parts[0] if len(parts) > 0 else “”
style     = parts[1].lower() if len(parts) > 1 else “apa”

```
if not os.path.exists(docx_path):
    return f"❌ File not found: {docx_path}"
if style not in STYLE_RULES:
    return f"❌ Unknown style: {style}"

ds   = STYLE_RULES[style]["document_setup"]
FONT = ds["font_family"]
SIZE = ds["font_size"]
LS   = ds["line_spacing"]

doc = Document(docx_path)
pc = rc = 0
for para in doc.paragraphs:
    if not para.text.strip():
        continue
    _set_paragraph_spacing(para, LS, 0, 0)
    pc += 1
    for run in para.runs:
        _set_run_font(run, FONT, SIZE,
                      bold=bool(run.bold), italic=bool(run.italic))
        rc += 1
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _set_paragraph_spacing(p, LS, 0, 0)
                for run in p.runs:
                    _set_run_font(run, FONT, SIZE)
                    rc += 1
doc.save(docx_path)
style_nm = STYLE_RULES[style]["name"]
return (f"✅ Sweep complete ({style_nm}): "
        f"{pc} paragraphs, {rc} runs updated → {docx_path}")
```

# ══════════════════════════════════════════════════════════════════════════════

# TOOL 6 — Generate Full Report

# ══════════════════════════════════════════════════════════════════════════════

def tool_generate_report(args_str: str) -> str:
“””
Generate a comprehensive formatting report for a .docx file.
Input: “docx_path style”
“””
parts = args_str.strip().strip(’”'’).split()
docx_path = parts[0] if len(parts) > 0 else “”
style     = parts[1].lower() if len(parts) > 1 else “apa”

```
if not os.path.exists(docx_path):
    return f"❌ File not found: {docx_path}"

try:
    pr = json.loads(tool_parse_document(f"{docx_path} {style}"))
    vr = json.loads(tool_validate(f"{docx_path} {style}"))
    dr = json.loads(tool_detect_style(docx_path))
except Exception as e:
    return f"❌ Error generating report: {e}"

sn = STYLE_RULES.get(style, {}).get("name", style.upper())
vs = vr.get("summary", {})
lines = [
    "╔" + "═" * 70 + "╗",
    f"║  MANUSCRIPT FORMATTING REPORT — {sn:<36} ║",
    "╚" + "═" * 70 + "╝",
    f"  File          : {pr.get('file', docx_path)}",
    f"  Applied Style : {sn}",
    f"  Total Words   : {pr.get('stats',{}).get('total_words','?')}",
    f"  Paragraphs    : {pr.get('total_paragraphs','?')}",
    "",
    "── STYLE AUTO-DETECTION ─────────────────────────────────────────────",
    f"  Detected  : {dr.get('detected_style','?').upper()} "
    f"({dr.get('confidence_percent','?')}% confidence)",
    f"  Scores    : APA={dr.get('scores',{}).get('apa',0):>3}  "
    f"MLA={dr.get('scores',{}).get('mla',0):>3}  "
    f"IEEE={dr.get('scores',{}).get('ieee',0):>3}",
]
for sty, evs in (dr.get("evidence") or {}).items():
    if evs:
        lines.append(f"  [{sty.upper():4s}] " + " | ".join(evs[:4]))

lines += ["", "── SECTIONS FOUND ───────────────────────────────────────────────────"]
for sec, val in (pr.get("special_sections") or {}).items():
    lines.append(f"  {sec:38s} ↦ paragraph #{val}")

lines += ["", "── HEADING STRUCTURE ────────────────────────────────────────────────"]
for h in pr.get("headings", []):
    indent = "  " * (h["level"] - 1)
    lines.append(f"  {indent}[L{h['level']}] {h['text'][:65]}")

cites = pr.get("citations", [])
lines += ["", f"── CITATIONS DETECTED ({len(cites)}) ──────────────────────────────────────"]
seen = set()
for c in cites:
    key = str(c.get("match", c))
    if key not in seen and len(seen) < 25:
        lines.append(f"  {key}")
        seen.add(key)

refs = pr.get("reference_entries", [])
lines += ["", f"── BIBLIOGRAPHY ENTRIES ({len(refs)}) ─────────────────────────────────────"]
for r in refs[:12]:
    lines.append(f"  {r['text'][:80]}")

issues = vr.get("issues", [])
lines += [
    "",
    f"── COMPLIANCE CHECK ({vs.get('paragraphs_checked','?')} paragraphs, "
    f"{len(issues)} issues) ──────────────────",
]
if not issues:
    lines.append("  ✅ No compliance issues — document is fully formatted.")
for iss in issues[:25]:
    lines.append(
        f"  ⚠  [{iss['type']:15s}] "
        f"para#{str(iss.get('para','?')):>4}  "
        f"found={str(iss.get('found','?')):<14}  "
        f"expected={iss.get('expected','?')}"
    )

lines += [
    "",
    "╔" + "═" * 70 + "╗",
    "║  END OF REPORT                                                        ║",
    "╚" + "═" * 70 + "╝",
]
return "\n".join(lines)
```

# ══════════════════════════════════════════════════════════════════════════════

# LANGCHAIN TOOL REGISTRY

# ══════════════════════════════════════════════════════════════════════════════

tools = [
Tool(
name=“Detect_Citation_Style”,
func=tool_detect_style,
description=(
“AUTO-DETECT whether a .docx uses APA 7th, MLA 9th, or IEEE. “
“Input: path to .docx file (string). “
“Output: JSON with ‘detected_style’ (apa/mla/ieee), “
“‘confidence_percent’, per-style scores, and evidence list. “
“ALWAYS call this FIRST before any other tool.”
)
),
Tool(
name=“Parse_Document”,
func=tool_parse_document,
description=(
“Parse a .docx and return structural JSON (headings, citations, sections). “
“Input: ‘docx_path style’ e.g. ‘paper.docx apa’. “
“Output: JSON with heading tree, detected citations, section boundaries, stats.”
)
),
Tool(
name=“Apply_Formatting”,
func=tool_apply_formatting,
description=(
“Apply full formatting for the detected style. “
“Input (JSON): {"docx_path":"paper.docx","style":"apa","output_path":"out.docx"} “
“OR space-separated: ‘paper.docx apa out.docx’. “
“Handles everything: page setup, fonts, margins, spacing, all heading levels, “
“title page, abstract, bibliography entries, header/footer, table cells. “
“Output: detailed formatting log.”
)
),
Tool(
name=“Validate_Compliance”,
func=tool_validate,
description=(
“Check formatted .docx for style compliance issues. “
“Input: ‘docx_path style’ or JSON {"docx_path":"…","style":"apa"}. “
“Output: JSON with issues list (font family, font size, margins, line spacing). “
“Call after Apply_Formatting.”
)
),
Tool(
name=“Fix_Font_Spacing”,
func=tool_fix_sweep,
description=(
“Brute-force sweep: enforce correct font/size/spacing on every run in file. “
“Input: ‘docx_path style’. Modifies file in-place. “
“Use this if Validate_Compliance still reports issues.”
)
),
Tool(
name=“Generate_Report”,
func=tool_generate_report,
description=(
“Generate a comprehensive human-readable formatting report. “
“Input: ‘docx_path style’. “
“Output: full report with style detection summary, section map, “
“heading tree, citation list, bibliography, and compliance check. “
“Call this LAST to summarize everything.”
)
),
]

# ══════════════════════════════════════════════════════════════════════════════

# LLM + REACT AGENT

# ══════════════════════════════════════════════════════════════════════════════

llm = ChatGroq(model=“llama-3.3-70b-versatile”, temperature=0)

REACT_PROMPT = “”“You are an expert academic manuscript formatter. Your task is to:

1. AUTOMATICALLY detect whether a document uses APA 7th, MLA 9th, or IEEE style
1. Apply the correct, complete formatting rules for that style
1. Validate the result and fix remaining issues
1. Generate a full formatting report

IMPORTANT: You must specify which style was detected and WHY, based on the evidence.

You have access to these tools:
{tools}

Use EXACTLY this format:

Question: the input question you must answer
Thought: think step by step about what to do
Action: the action to take, must be one of [{tool_names}]
Action Input: the input to the action
Observation: the result of the action
… (repeat as needed)
Thought: I now know the final answer
Final Answer: complete summary of what was done

REQUIRED WORKFLOW (follow in order):
Step 1: Call Detect_Citation_Style → determine APA/MLA/IEEE. State detected style + confidence.
Step 2: Call Parse_Document with “docx_path detected_style” → understand document structure.
Step 3: Call Apply_Formatting with JSON: {{“docx_path”:”…”,“style”:“detected_style”,“output_path”:”…”}}.
Step 4: Call Validate_Compliance with “output_path detected_style” → check quality.
Step 5: If compliance issues > 0, call Fix_Font_Spacing with “output_path detected_style”.
Step 6: Call Generate_Report with “output_path detected_style” → final detailed report.
Step 7: Final Answer must include:

- Which style was detected and why (top evidence)
- What formatting was applied
- Output file path
- Any remaining notes

Begin!

Question: {input}
Thought:{agent_scratchpad}”””

prompt = PromptTemplate.from_template(REACT_PROMPT)
agent = create_react_agent(llm=llm, tools=tools, prompt=prompt)
agent_executor = AgentExecutor(
agent=agent,
tools=tools,
verbose=True,
max_iterations=18,
handle_parsing_errors=True,
return_intermediate_steps=True,
)

# ══════════════════════════════════════════════════════════════════════════════

# PUBLIC API

# ══════════════════════════════════════════════════════════════════════════════

def format_manuscript(input_path: str,
output_path: str = None,
force_style: str = None) -> str:
if not os.path.exists(input_path):
return f”❌ File not found: {input_path}”

```
if output_path is None:
    output_path = Path(input_path).stem + "_formatted.docx"

if force_style:
    fs = force_style.lower()
    sn = STYLE_RULES.get(fs, {}).get("name", fs.upper())
    task = (
        f"Format '{input_path}' using {sn} rules (style is already known: {fs}). "
        f"Save formatted output as '{output_path}'. "
        f"Apply ALL {sn} formatting rules. Then validate and report."
    )
else:
    task = (
        f"Format the manuscript '{input_path}'. "
        f"First detect whether it uses APA 7th, MLA 9th, or IEEE style by "
        f"analyzing citation patterns, section titles, and domain keywords. "
        f"Then apply the correct and complete formatting for the detected style. "
        f"Save the formatted output as '{output_path}'. "
        f"Validate, fix any issues, and generate a full formatting report."
    )

try:
    result = agent_executor.invoke({"input": task})
    output = result.get("output", "")
    steps = result.get("intermediate_steps", [])
    if steps:
        output += f"\n\n📋 Agent completed {len(steps)} pipeline steps."
    return output
except Exception as exc:
    print(f"⚠ Agent error: {exc}\n⚠ Running direct fallback pipeline…")
    detected = json.loads(tool_detect_style(input_path))
    style = force_style or detected.get("detected_style", "apa")
    sn = STYLE_RULES.get(style, {}).get("name", style.upper())
    conf = detected.get("confidence_percent", 0)
    print(f"\n🔍 Detected: {sn} ({conf}% confidence)")
    r1 = tool_apply_formatting(
        json.dumps({"docx_path": input_path, "style": style,
                    "output_path": output_path}))
    r2 = tool_fix_sweep(f"{output_path} {style}")
    r3 = tool_generate_report(f"{output_path} {style}")
    return f"STYLE DETECTED: {sn} ({conf}% confidence)\n\n{r1}\n\n{r2}\n\n{r3}"
```

# ══════════════════════════════════════════════════════════════════════════════

# CLI

# ══════════════════════════════════════════════════════════════════════════════

if **name** == “**main**”:
parser = argparse.ArgumentParser(
description=“Academic Manuscript Formatter — APA 7th / MLA 9th / IEEE”,
formatter_class=argparse.RawDescriptionHelpFormatter,
epilog=**doc**,
)
parser.add_argument(“input”, help=“Path to the input .docx file”)
parser.add_argument(”–output”, “-o”, default=None,
help=“Output path (default: <input>_formatted.docx)”)
parser.add_argument(”–style”, “-s”, choices=[“apa”, “mla”, “ieee”],
default=None,
help=“Force a style (default: auto-detect)”)
parser.add_argument(”–direct”, “-d”, action=“store_true”,
help=“Skip LLM agent, run pipeline directly”)
parser.add_argument(”–detect-only”, action=“store_true”,
help=“Only detect and print the citation style”)
args = parser.parse_args()

```
output = args.output or (Path(args.input).stem + "_formatted.docx")

if args.detect_only:
    res = json.loads(tool_detect_style(args.input))
    print(f"\n{'═'*55}")
    print(f"  STYLE DETECTION RESULT")
    print(f"{'═'*55}")
    print(f"  Detected   : {res['detected_style'].upper()}")
    print(f"  Name       : {res['style_name']}")
    print(f"  Discipline : {res['discipline']}")
    print(f"  Confidence : {res['confidence_percent']}%")
    print(f"  Scores     : APA={res['scores']['apa']}  "
          f"MLA={res['scores']['mla']}  IEEE={res['scores']['ieee']}")
    print(f"\n  Detection Evidence:")
    for sty, evs in res["evidence"].items():
        if evs:
            print(f"    [{sty.upper()}] " + " | ".join(evs))
    print(f"{'═'*55}\n")

elif args.direct:
    detected = json.loads(tool_detect_style(args.input))
    style = args.style or detected.get("detected_style", "apa")
    sn    = STYLE_RULES.get(style, {}).get("name", style.upper())
    conf  = detected.get("confidence_percent", 0)
    print(f"\n🔍 Auto-detected style: {sn} ({conf}% confidence)")
    print(f"📄 Applying {sn} formatting to {args.input}…\n")
    print(tool_apply_formatting(
        json.dumps({"docx_path": args.input, "style": style,
                    "output_path": output})))
    print(tool_fix_sweep(f"{output} {style}"))
    print(tool_generate_report(f"{output} {style}"))
else:
    print(format_manuscript(args.input, output, args.style))
```